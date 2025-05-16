from django.contrib import messages
from django.core.exceptions import ObjectDoesNotExist
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.mixins import LoginRequiredMixin
from .models import *
from django.views.generic import ListView, DetailView, View, UpdateView
from django.shortcuts import render, get_object_or_404, redirect
from django.utils import timezone
from .forms import *
from django.core.mail import EmailMessage,send_mail
from django.contrib.auth.models import User
from django.shortcuts import reverse, HttpResponse
from django.db.models import Q, F, Case, When, IntegerField, Value
from django.db.models.functions import Coalesce,Cast
from django.template.loader import render_to_string, get_template
from django.conf import settings
from xhtml2pdf import pisa
from io import BytesIO
from datetime import datetime
from django.contrib.admin.views.decorators import staff_member_required
from django.views.decorators.csrf import csrf_exempt
from itertools import chain
import qrcode
import base64
import os
from datetime import date
from django.http import HttpResponse
from django.conf import settings
from docx import Document
import boto3
from django.http import JsonResponse
import json
from django.utils.html import escape
from django.core.serializers.json import DjangoJSONEncoder
import json

email_master = settings.EMAIL_MASTER


#CRM

@staff_member_required
def update_lieferanten_status(request, pk):
    bestellung = get_object_or_404(LieferantenBestellungen, id=pk)
    
    if request.method == "POST":
        form = LieferantenStatusForm(request.POST)
        if form.is_valid():
            new_status = form.save(commit=False)
            new_status.date = now()
            new_status.order = bestellung
            new_status.save()
            bestellung.status.add(new_status)
            messages.success(request, "Status erfolgreich aktualisiert.")
            return redirect('store:lieferanten_bestellungen')
    else:
        form = LieferantenStatusForm()
    
    return render(request, 'crm/lieferant-update-status.html', {'form': form, 'bestellung': bestellung})


@staff_member_required
def lieferanten_bestellungen(request):
    bestellungen = LieferantenBestellungen.objects.select_related('auftrag').prefetch_related(
        'artikel_bestellungen',
        'auftrag__elementeitems_bestellung__element_nr',
        'auftrag__elementeitems_bestellung__artikel',
    ).order_by('-start_date')

    # Dictionary: bestellungs_id → Liste der Cart-Items
    elemente_data = {}

    for bestellung in bestellungen:
        if bestellung.auftrag:
            elemente_items = bestellung.auftrag.elementeitems_bestellung.all()
            elemente_data[bestellung.id] = [
                {
                    "element_nr": item.element_nr.elementnr if item.element_nr else "–",
                    "anzahl": item.anzahl,
                    "artikel": item.artikel.artikelnr if item.artikel else "–",
                    "masse": f"{item.element_nr.aussenbreite}mm x {item.element_nr.aussenhöhe}mm" if item.element_nr else "–"
                }
                for item in elemente_items
            ]

    return render(request, 'crm/lieferanten-bestellungen.html', {
        'lieferanten_bestellungen': bestellungen,
        'elemente_data_json': json.dumps(elemente_data, cls=DjangoJSONEncoder)
    })


@staff_member_required
def lieferant_send_order_email(request, pk):
    artikel = get_object_or_404(Artikel, pk=pk)

    if request.method == "POST":
        anzahl = request.POST.get("anzahl", 1)  # Default to 1 if not provided
        lieferant_id = request.POST.get("lieferant")  # Get the selected Lieferant dynamically

        # Validate and retrieve the Lieferant
        lieferant = None
        if lieferant_id:
            try:
                lieferant = get_object_or_404(Lieferanten, pk=int(lieferant_id))
            except (ValueError, TypeError):
                messages.error(request, "Ungültige Lieferant-ID.")
                return redirect('store:crm_artikel')

        # Fallback to the article's default Lieferant
        if not lieferant:
            lieferant = artikel.lieferant

        if lieferant and lieferant.email:
            bestellung = LieferantenBestellungen.objects.create(lieferant_id=lieferant.id,)
            bestellartikel = LieferantenBestellungenArtikel.objects.create(order=bestellung, artikel=artikel, anzahl=anzahl)
            status = LieferantenStatus.objects.create(order=bestellung, name="In Bearbeitung")
            bestellung.status.add(status)
            
            # Prepare email content
            subject = f"Bestellung für Artikel {artikel.artikelnr}"
            template = render_to_string('emails/order_email.html', {
                'lieferant': lieferant,
                'artikel': artikel,
                'anzahl': anzahl,
            })

            # Send email
            email = EmailMessage(
                subject,
                template,
                settings.EMAIL_HOST_USER,  # Sender email
                [lieferant.email],  # Recipients
                bcc=email_master
            )
            email.content_subtype = 'html'  # Send as HTML
            email.send()

            messages.success(
                request,
                f"Die Bestellung wurde erfolgreich an {lieferant.name} ({lieferant.email}) gesendet."
            )
        else:
            messages.error(request, "Keine gültige E-Mail-Adresse für diesen Lieferanten verfügbar.")

    return redirect('store:lieferanten_bestellungen')


@staff_member_required
def elemente_bestellung_detail(request, pk, betrieb):
    bestellung = get_object_or_404(Elemente_Bestellungen, id=pk)

    # Get all cart items for the order
    cart_items = ElementeCartItem.objects.filter(order=bestellung).select_related("artikel")

    # Fetch all Lieferanten (suppliers)
    lieferanten = Lieferanten.objects.all()

    kunde = Kunde.objects.get(firmenname=betrieb)

    def get_dichtungstyp(element_input, kunde=None):
        if isinstance(element_input, Elemente):
            element = element_input
        else:
            try:
                element_nr = int(element_input)
                queryset = Elemente.objects.filter(elementnr=element_nr)
                if kunde:
                    queryset = queryset.filter(kunde=kunde)
                element = queryset.first()
            except (ValueError, TypeError):
                return "Unbekannt"

        return element.bemerkung if element and element.bemerkung else "Unbekannt"

    # Create list of elements with additional details from `ElementeCartItem` model
    elemente_list = []
    show_lieferantenartikel = False

    for item in cart_items:
        artikel = item.artikel
        dichtungstyp = get_dichtungstyp(item.element_nr_id)

        artikel_data = {
            "id": artikel.id if artikel else None,
            "artikelnr": artikel.artikelnr if artikel else "Unbekannt",
            "name": artikel.name if artikel else "Unbekannt",
            "aussenbreite": artikel.aussenbreite if artikel else "Unbekannt",
            "aussenhöhe": artikel.aussenhöhe if artikel else "Unbekannt",
            "lieferant": artikel.lieferant if artikel and artikel.lieferant else None,
            "zubehoerartikelnr": artikel.zubehoerartikelnr if artikel else None,
            "lieferantenartikel": artikel.lieferantenartikel if artikel else None,
        } if artikel else None

        if artikel and artikel.lieferantenartikel:
            show_lieferantenartikel = True

        cart_items_nr = cart_items.get(order=bestellung,element_nr=item.element_nr_id)

        elemente_list.append({
            'id': item.id,
            "element_nr": cart_items_nr.element_nr,
            "dichtungstyp": dichtungstyp,
            "artikel": artikel_data,
            "masse": f"{artikel.aussenbreite}mm x {artikel.aussenhöhe}mm" if artikel else "Unbekannt",
            "stk_zahl": item.anzahl
        })

    error_message = None

    # ✅ Handle form submission
    if request.method == "POST":
        if "send_order" in request.POST:
            # ✅ Send Order Logic (Restoring Email Functionality)
            bestellung_artikel_list = []
            lieferant = None

            for item in cart_items:
                artikel = item.artikel
                if artikel:
                    if not lieferant:
                        lieferant = artikel.lieferant

                    bestellung_artikel_list.append({
                        "artikel": artikel,
                        "anzahl": item.anzahl
                    })

            if lieferant and bestellung_artikel_list:
                try:
                    lieferanten_bestellung = LieferantenBestellungen.objects.create(lieferant=lieferant, auftrag=bestellung)
                    status = LieferantenStatus.objects.create(order=lieferanten_bestellung, name="Versendet")
                    lieferanten_bestellung.status.set([status])

                    if bestellung.wer == "lager":
                        bestellung.wer = "lieferant"
                        bestellung.save()

                    for bestell_artikel in bestellung_artikel_list:
                        LieferantenBestellungenArtikel.objects.create(
                            order=lieferanten_bestellung,
                            artikel=bestell_artikel["artikel"],
                            anzahl=bestell_artikel["anzahl"]
                        )

                    bestellung.status = "Versendet"
                    bestellung.wer = "Lieferant"
                    bestellung.save()

                    # ✅ Send Email
                    if lieferant.email:
                        subject = f"Bestellung für {bestellung.kunden_nr} - Auftrag #{bestellung.id}"
                        template = render_to_string('emails/bestellung_lieferant.html', {
                            'lieferant': lieferant,
                            'artikel_liste': bestellung_artikel_list,
                            "show_lieferantenartikel": show_lieferantenartikel,
                        })

                        email = EmailMessage(
                            subject,
                            template,
                            settings.EMAIL_HOST_USER,
                            [lieferant.email],  
                            bcc=email_master
                        )
                        email.content_subtype = 'html'
                        email.send()

                        messages.success(request, f"Bestellung erfolgreich versendet und E-Mail an {lieferant.name} gesendet.")
                    else:
                        messages.error(request, "Keine gültige E-Mail-Adresse für diesen Lieferanten verfügbar.")

                    return redirect("store:elemente_bestellung_detail", pk=pk, betrieb=betrieb)

                except Exception as e:
                    messages.error(request, f"Fehler beim Speichern der Bestellung: {str(e)}")
                    return redirect("store:elemente_bestellung_detail", pk=pk, betrieb=betrieb)

            else:
                messages.error(request, "Es wurden keine Artikel zur Bestellung hinzugefügt.")

        elif "element_nr" in request.POST and "anzahl" in request.POST:
            # ✅ Position hinzufügen (Update wenn vorhanden)
            element_nr = request.POST["element_nr"]
            anzahl = int(request.POST["anzahl"])

           

            # Check if element exists in `Elemente` model
            try:
                element = Elemente.objects.get(elementnr=element_nr, kunde=kunde)
            except Elemente.DoesNotExist:
                messages.error(request, "Dieses Element existiert nicht für den Kunden.")
                return redirect("store:elemente_bestellung_detail", pk=pk, betrieb=betrieb)
            except Elemente.MultipleObjectsReturned:
                messages.error(request, "Mehrere Elemente mit derselben Nummer gefunden – bitte prüfen Sie die Daten.")
                return redirect("store:elemente_bestellung_detail", pk=pk, betrieb=betrieb)

            if not element:
                error_message = "Das eingegebene Element existiert nicht."
            else:
                artikel = element.artikel
                existing_item = ElementeCartItem.objects.filter(order=bestellung, element_nr=element.id).first()

                if existing_item:
                    existing_item.anzahl += anzahl
                    existing_item.save()
                    messages.success(request, "Menge erfolgreich aktualisiert.")
                else:    
                    ElementeCartItem.objects.create(order=bestellung, element_nr=element, artikel=artikel, anzahl=anzahl)
                    messages.success(request, "Position erfolgreich hinzugefügt.")

                return redirect("store:elemente_bestellung_detail", pk=pk, betrieb=betrieb)

        elif "update_notizfeld" in request.POST:
            notiz = request.POST.get("notizfeld", "")
            bestellung.notizfeld = notiz
            bestellung.save()
            messages.success(request, "Notizfeld wurde aktualisiert.")
            return redirect("store:elemente_bestellung_detail", bestellung.id, betrieb)

    # ✅ Context for the template
    context = {
        "bestellung": {
            "id": bestellung.id,
            "kunden_nr": bestellung.kunden_nr,
            "status": bestellung.status,
            "start_date": bestellung.start_date,
            "montage": bestellung.montage,
            "notizfeld": bestellung.notizfeld
        },
        "elemente": elemente_list,
        "lieferanten": lieferanten,
        "betrieb": betrieb,
        "show_lieferantenartikel": show_lieferantenartikel,
        "error_message": error_message,
        "cart_items":cart_items,
    }

    return render(request, "crm/cms-elemente-bestellungen-detail.html", context)


@staff_member_required
def bestellung_elemente_detail_delete(request, pk, bestellung_id):
    # ✅ Bestellung & Kunde laden
    bestellung = get_object_or_404(Elemente_Bestellungen, id=bestellung_id)
    kunde = get_object_or_404(Kunde, interne_nummer=bestellung.kunden_nr)
    betrieb = kunde.firmenname

    # ✅ Zu löschendes Item holen
    item = ElementeCartItem.objects.filter(id=pk, order=bestellung).first()

    if not item:
        messages.error(request, "Diese Position existiert nicht oder gehört nicht zur Bestellung.")
        return redirect("store:elemente_bestellung_detail", pk=bestellung.id, betrieb=betrieb)

    # ✅ Löschen
    item.delete()
    messages.info(request, "Die Position wurde gelöscht.")
    return redirect("store:elemente_bestellung_detail", pk=bestellung.id, betrieb=betrieb)


@staff_member_required
def elemente_bestellung_detail_edit(request, pk, bestellung_id):
    bestellung = get_object_or_404(Elemente_Bestellungen, id=bestellung_id)
    kunde = get_object_or_404(Kunde, interne_nummer=bestellung.kunden_nr)
    betrieb = kunde.firmenname
    item = get_object_or_404(ElementeCartItem, order=bestellung, id=pk)

    if request.method == "POST":
        form = ElementeCartItemEditForm(request.POST)
        if form.is_valid():
            item.anzahl = form.cleaned_data["anzahl"]
            item.save()
            messages.success(request, "Element erfolgreich aktualisiert.")
            return redirect("store:elemente_bestellung_detail", pk=bestellung.id, betrieb=betrieb)
        else:
            messages.error(request, "Ungültige Eingabe. Bitte überprüfen Sie die Werte.")
    else:
        form = ElementeCartItemEditForm(initial={"anzahl": item.anzahl})

    context = {"form": form, "item": item, "bestellung": bestellung, "betrieb": betrieb}
    return render(request, "crm/crm-elemente-bestellung-edit.html", context)


@staff_member_required
def elemente_bestellung_delete(request, pk):
    bestellung = get_object_or_404(Elemente_Bestellungen, pk=pk)
    bestellung.delete()
    messages.success(request, f"Bestellung #{pk} wurde gelöscht.")
    return redirect('store:elemente_bestellungen')


@staff_member_required
def elemente_bestellungen(request):
    search_query = request.GET.get('search', '')

    if search_query:
        bestellungen = Elemente_Bestellungen.objects.filter(
            Q(id__icontains=search_query)
        ).order_by('-start_date')
    else:
        bestellungen = Elemente_Bestellungen.objects.all().order_by('-start_date')

    # Fetch Kunde details for each Bestellung
    bestellungen_list = []
    for bestellung in bestellungen:
        kunde = Kunde.objects.filter(interne_nummer=bestellung.kunden_nr).first()
        if not kunde:
            continue
        bestellungen_list.append({
            "id": bestellung.id,
            "kunden_nr": bestellung.kunden_nr,
            "status": bestellung.status,
            "start_date": bestellung.start_date,
            "montage": bestellung.montage,
            "betrieb_person": kunde.firmenname,
            "kunde": kunde,  # ➕ Das ist wichtig für das Modal
            "wer": bestellung.wer,
            "wer_display": bestellung.get_wer_display(),  # <--- das hier
        })

    context = {
        "bestellungen": bestellungen_list,
    }

    return render(request, "crm/cms-elemente-bestellungen.html", context)



def bestellformular_cart(request):
    """
    Handles the cart functionality, validates element-nr, saves the linked article,
    sends an email for checkout, and renders the bestellformular template.
    """
    if "kunden_nr" not in request.session:
        request.session["kunden_nr"] = None

    kunden_nr_form = KundenNrForm()
    form = ElementeCartItemForm()
    kunde_details = None
    cart_items_with_details = []

    # ✅ Ensure `kunden_nr` is valid
    kunden_nr = request.session.get("kunden_nr")
    if kunden_nr:
        kunde_details = Kunde.objects.filter(interne_nummer=kunden_nr).first()
        if not kunde_details:
            request.session["kunden_nr"] = None  # Reset if invalid

    if request.method == "POST":
        # **Step 1: Set Kunden-Nr.**
        if "set_kunden_nr" in request.POST:
            kunden_nr_form = KundenNrForm(request.POST)
            if kunden_nr_form.is_valid():
                kunden_nr = kunden_nr_form.cleaned_data["kunden_nr"]
                kunde_details = Kunde.objects.filter(interne_nummer=kunden_nr).first()

                if not kunde_details:
                    messages.warning(request, "Dieser Kunde existiert nicht. Bitte geben Sie eine gültige Kunden-Nr. ein.")
                else:
                    request.session["kunden_nr"] = kunden_nr
                    request.session.modified = True
                    return redirect("store:bestellformular_cart")

        # **Step 2: Reset Kunden-Nr.**
        elif "reset_kunden_nr" in request.POST:
            request.session["kunden_nr"] = None
            request.session.modified = True
            messages.success(request, "Kunden-Nr. wurde erfolgreich zurückgesetzt.")
            return redirect("store:bestellformular_cart")

        # **Step 3: Add Item to Cart**
        elif "add_to_cart" in request.POST:
            form = ElementeCartItemForm(request.POST)
            if form.is_valid():
                kunden_nr = request.POST.get("kunden_nr") or request.session.get("kunden_nr")

                if not kunden_nr:
                    messages.warning(request, "Bitte wählen Sie zuerst einen Kunden aus.")
                    return redirect("store:bestellformular_cart")

                kunde_details = get_object_or_404(Kunde, interne_nummer=kunden_nr)

                element_nr = form.cleaned_data["element_nr"]
                anzahl = form.cleaned_data["anzahl"]

                # ✅ Check if the element exists and is linked to an article
                element = Elemente.objects.filter(elementnr=element_nr, kunde__interne_nummer=kunden_nr).select_related("artikel").first()
                if not element:
                    messages.warning(request, f"Das Element-Nr. '{element_nr}' existiert nicht für diesen Kunden.")
                else:
                    artikel = element.artikel if element.artikel else None

                    # ✅ Ensure an order exists
                    order, created = Elemente_Bestellungen.objects.get_or_create(
                        kunden_nr=str(kunden_nr), status="offen", wer="lager"
                    )

                    if created:
                        order.save()

                    # ✅ Check if the element is already in the cart
                    # existierendes Element korrekt prüfen
                    existing_item = ElementeCartItem.objects.filter(order=order, element_nr=element).first()


                    if existing_item:
                        existing_item.anzahl += anzahl
                        existing_item.save()
                        messages.success(request, f"Die Anzahl für Element-Nr. '{element_nr}' wurde erhöht.")
                    else:
                        ElementeCartItem.objects.create(
                            order=order,
                            element_nr=element,
                            artikel=artikel,  # ✅ Save article
                            anzahl=anzahl,
                        )
                        messages.success(request, f"Element-Nr. '{element_nr}' erfolgreich hinzugefügt.")

                    return redirect("store:bestellformular_cart")

        # **Step 4: Delete Item from Cart**
        elif "delete_item" in request.POST:
            delete_item_id = request.POST.get("delete_item_id")
            try:
                item_to_delete = ElementeCartItem.objects.get(id=delete_item_id)
                item_to_delete.delete()
                messages.success(request, "Artikel erfolgreich aus dem Warenkorb entfernt.")
            except ElementeCartItem.DoesNotExist:
                messages.error(request, "Artikel konnte nicht gefunden werden.")
            return redirect("store:bestellformular_cart")

        # **Step 5: Checkout - Save Montage & Send Email**
        elif "checkout" in request.POST:
            montage = request.POST.get("montage", "Nein")  # ✅ Ensure montage is retrieved from form

            order = Elemente_Bestellungen.objects.filter(kunden_nr=str(kunden_nr), status="offen").first()
            
            if not order or not order.elementeitems_bestellung.exists():
                messages.warning(request, "Bestellung konnte nicht abgeschlossen werden. Bitte fügen Sie Artikel hinzu.")
                return redirect("store:bestellformular_cart")

            # ✅ Update `montage` and set status to "teilweise"
            order.montage = montage
            order.status = "offen"
            order.save(update_fields=["montage", "status"])  # ✅ Ensure fields are stored in the database

            # ✅ Fetch customer details
            kunde_details = Kunde.objects.filter(interne_nummer=kunden_nr).first()
            if not kunde_details:
                messages.warning(request, "Kundendetails nicht verfügbar. Bitte erneut versuchen.")
                return redirect("store:bestellformular_cart")

            betrieb_person = kunde_details.firmenname if kunde_details.firmenname else f"{kunde_details.vorname} {kunde_details.nachname}"

            crm_address = kunde_details.kunde_address.first()
            adresse = f"{crm_address.crm_strasse} {crm_address.crm_nr}" if crm_address else "Adresse nicht verfügbar"
            plz = crm_address.crm_plz if crm_address else "PLZ nicht verfügbar"
            ort = crm_address.crm_ort if crm_address else "Ort nicht verfügbar"

            elemente_list = list(order.elementeitems_bestellung.all())  # Convert QuerySet to list

            # ✅ Send Email
            subject = f"Elemente Bestellung-Nr. {order.id}, Kunden-Nr. {kunden_nr} {betrieb_person}"
            template = render_to_string("crm/mail-bestellung-elemente.html", {
                "kunden_nr": kunden_nr,
                "id": order.id,
                "betrieb_person": betrieb_person,
                "adresse": adresse,
                "plz": plz,
                "ort": ort,
                "elemente_list": elemente_list,
                "montage": montage,  # ✅ Pass Montage correctly to email
            })

            try:
                email = EmailMessage(subject, template, settings.EMAIL_HOST_USER, email_master)
                email.content_subtype = "html"
                email.send()
                messages.success(request, "Bestellung erfolgreich abgeschlossen.")
            except Exception as e:
                messages.error(request, f"E-Mail konnte nicht versendet werden: {str(e)}")

            # ✅ Second Email Confirmation to Kunde
            subject = f"Auftragsbestätigung Kühlschrankdichtungen - Bestellung {order.id}"
            template = render_to_string("crm/mail-bestaetigung-elemente.html", {
                "id": order.id,
                "date": order.start_date,
                "kunden_nr": kunden_nr,
                "betrieb_person": betrieb_person,
                "montage": montage,  # ✅ Ensure montage is correctly passed
            })

            try:
                email = EmailMessage(subject, template, settings.EMAIL_HOST_USER, [kunde_details.email])
                email.content_subtype = "html"
                email.send()
            except Exception as e:
                messages.error(request, f"Fehler beim Senden der E-Mail: {str(e)}")

            # ✅ Exclude items from the next cart view
            request.session["kunden_nr"] = None
            request.session.modified = True
            return redirect("store:bestellformular_cart")


    # ✅ Retrieve cart contents and exclude completed orders
    if kunden_nr:
        order = Elemente_Bestellungen.objects.filter(kunden_nr=kunden_nr, status="offen").first()

        # Ensure only items from open orders are shown
        if order and order.elementeitems_bestellung.exists():
            cart_items_with_details = [
                {
                    "element_nr": item.element_nr,
                    "anzahl": item.anzahl,
                    "artikel": item.artikel,
                    "kuehlposition": item.artikel.artikel_elemente.first().kuehlposition if item.artikel and item.artikel.artikel_elemente.exists() else 'N/A',
                    "bezeichnung": item.artikel.artikel_elemente.first().bezeichnung if item.artikel and item.artikel.artikel_elemente.exists() else 'N/A',
                    "id": item.id,
                }
                for item in order.elementeitems_bestellung.all()
            ]
        else:
            cart_items_with_details = []


    return render(request, "bestellformular-neu.html", {
        "kunden_nr_form": kunden_nr_form,
        "form": form,
        "cart_items_with_details": cart_items_with_details,
        "kunden_nr": kunden_nr,
        "kunde_details": kunde_details,
    })
# end new warenkorb elemente



# Start CRM Artikel
@staff_member_required
def get_lieferant_email(request, lieferant_id):
    try:
        lieferant = Lieferanten.objects.get(pk=lieferant_id)
        return JsonResponse({'email': lieferant.email})
    except Lieferanten.DoesNotExist:
        return JsonResponse({'email': None})


@staff_member_required
def crm_artikel_nettopreis_edit(request, pk):
    artikel = get_object_or_404(Artikel, pk=pk)
    if request.method == 'POST':
        form = NettopreisArtikelForm(request.POST, instance=artikel)
        if form.is_valid():
            form.save()
            messages.success(request, "Einkaufspreis erfolgreich aktualisiert.")
            return redirect('store:crm_artikel')  # Redirect back to Artikel list
    else:
        form = NettopreisArtikelForm(instance=artikel)
    return render(request, 'crm/crm-artikel-nettopreis-edit.html', {'form': form, 'artikel': artikel})

@staff_member_required
def crm_artikel_preiscode_edit(request, pk):
    artikel = get_object_or_404(Artikel, pk=pk)
    if request.method == 'POST':
        form = PreiscodeArtikelForm(request.POST, instance=artikel)
        if form.is_valid():
            form.save()
            messages.success(request, "Preiscode erfolgreich aktualisiert.")
            return redirect('store:crm_artikel')  # Redirect back to the Artikel list
    else:
        form = PreiscodeArtikelForm(instance=artikel)
    return render(request, 'crm/crm-artikel-preiscode-edit.html', {'form': form, 'artikel': artikel})

# Edit Lagerbestand View
@staff_member_required
def crm_artikel_lagerbestand_edit(request, pk):
    artikel = get_object_or_404(Artikel, pk=pk)
    if request.method == 'POST':
        form = LagerbestandForm(request.POST, instance=artikel)
        if form.is_valid():
            form.save()
            messages.success(request, "Lagerbestand erfolgreich aktualisiert.")
            return redirect('store:crm_artikel')
    else:
        form = LagerbestandForm(instance=artikel)
    return render(request, 'crm/crm-artikel-lagerbestand-bearbeiten.html', {'form': form, 'artikel': artikel})

# Edit Lagerort View
@staff_member_required
def crm_artikel_lagerort_edit(request, pk):
    artikel = get_object_or_404(Artikel, pk=pk)
    if request.method == 'POST':
        form = LagerortForm(request.POST, instance=artikel)
        if form.is_valid():
            form.save()
            messages.success(request, "Lagerort erfolgreich aktualisiert.")
            return redirect('store:crm_artikel')
    else:
        form = LagerortForm(instance=artikel)
    return render(request, 'crm/crm-artikel-lagerort-bearbeiten.html', {'form': form, 'artikel': artikel})

# View to list all Preiscodes
@staff_member_required
def crm_preiscode(request):
    search_query = request.GET.get('search', '')
    if search_query:
        preiscodes = Preiscode.objects.filter(
            Q(preiscode__icontains=search_query) |
            Q(faktor__icontains=search_query) |
            Q(transportkosten__icontains=search_query) |
            Q(rabatt__icontains=search_query) |
            Q(preisanpassung__icontains=search_query)
        ).order_by('-id')
    else:
        preiscodes = Preiscode.objects.all().order_by('preiscode')
    return render(request, 'crm/crm-preiscodes.html', {'preiscodes': preiscodes})

# View to create a new Preiscode
@staff_member_required
def crm_preiscode_create(request):
    if request.method == 'POST':
        form = PreiscodeForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('store:crm_preiscode')
    else:
        form = PreiscodeForm()
    return render(request, 'crm/crm-preiscode-erfassen.html', {'form': form})

# View to edit an existing Preiscode
@staff_member_required
def crm_preiscode_edit(request, pk):
    preiscode = get_object_or_404(Preiscode, pk=pk)
    if request.method == 'POST':
        form = PreiscodeForm(request.POST, instance=preiscode)
        if form.is_valid():
            form.save()
            return redirect('store:crm_preiscode')
    else:
        form = PreiscodeForm(instance=preiscode)
    return render(request, 'crm/crm-preiscode-bearbeiten.html', {'form': form})

# View to delete an existing Preiscode
@staff_member_required
def crm_preiscode_delete(request, pk):
    preiscode = get_object_or_404(Preiscode, pk=pk)
    preiscode.delete()
    messages.info(request, "Der Preiscode wurde gelöscht.")
    return redirect('store:crm_preiscode')

@staff_member_required
def change_artikel_lieferant(request, pk):
    artikel = get_object_or_404(Artikel, pk=pk)
    if request.method == 'POST':
        lieferant_id = request.POST.get('lieferant')
        lieferant = get_object_or_404(Lieferanten, pk=lieferant_id)
        artikel.lieferant = lieferant
        artikel.save()
        messages.success(request, f"Lieferant für Artikel {artikel.artikelnr} wurde erfolgreich geändert.")
    return redirect('store:crm_artikel')  # Redirect to the Artikel list


@staff_member_required
def crm_artikel(request):
    search_query = request.GET.get('search', '')

    # Annotate with a numeric field, defaulting to 0 if not numeric
    artikel_queryset = Artikel.objects.annotate(
        artikelnr_numeric=Case(
            When(artikelnr__regex=r'^\d+$', then=Cast('artikelnr', IntegerField())),
            default=Value(0),  # Default for non-numeric artikelnr
            output_field=IntegerField()
        )
    )

    if search_query:
        # Filter by search query
        artikel = artikel_queryset.filter(
            Q(artikelnr__icontains=search_query) |
            Q(name__icontains=search_query)
        ).order_by('artikelnr_numeric', 'artikelnr')  # Sort by numeric prefix first
    else:
        # Default queryset ordering
        artikel = artikel_queryset.order_by('artikelnr_numeric', 'artikelnr')

    lieferanten = Lieferanten.objects.all()
    return render(request, 'crm/crm-artikel.html', {'artikel': artikel, 'lieferanten': lieferanten})


# View to create a new Artikel
@staff_member_required
def crm_artikel_create(request):
    if request.method == 'POST':
        form = ArtikelForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('store:crm_artikel')
    else:
        form = ArtikelForm()

    return render(request, 'crm/crm-artikel-erfassen.html',{'form': form,})

# View to edit an existing Artikel
@staff_member_required
def crm_artikel_edit(request, pk):
    artikel = get_object_or_404(Artikel, pk=pk)
    if request.method == 'POST':
        form = ArtikelForm(request.POST, instance=artikel)
        if form.is_valid():
            form.save()
            return redirect('store:crm_artikel')
        else:
            print(form.errors)  # Debug: Print form errors
    else:
        form = ArtikelForm(instance=artikel)
    return render(request, 'crm/crm-artikel-bearbeiten.html', {'form': form, 'artikel': artikel})


@staff_member_required
def crm_artikel_delete(request, pk):
    eintrag = get_object_or_404(Artikel, pk=pk)
    eintrag.delete()
    messages.info(request, "Der Artikel wurde gelöscht.")
    return redirect("store:crm_artikel")

@staff_member_required
def crm_lagerbestand(request):
    # Get the search parameters from GET
    dichtungstyp = request.GET.get('dichtungstyp', '')
    aussenbreite = request.GET.get('aussenbreite', '')
    aussenhoehe = request.GET.get('aussenhöhe', '')  # Using aussenhöhe with umlaut
    marke = request.GET.get('marke', '')

    # Start with the base queryset
    lager_queryset = CRMLager.objects.all()

    # Apply filters for each search field
    if dichtungstyp:
        lager_queryset = lager_queryset.filter(dichtungstyp__icontains=dichtungstyp)

    # Filter with tolerance for aussenbreite (+/- 10)
    if aussenbreite:
        try:
            aussenbreite_value = int(aussenbreite)
            lager_queryset = lager_queryset.filter(
                aussenbreite__gte=aussenbreite_value - 10,
                aussenbreite__lte=aussenbreite_value + 10
            )
        except ValueError:
            pass  # In case a non-integer is entered

    # Filter with tolerance for aussenhöhe (+/- 10)
    if aussenhoehe:
        try:
            aussenhoehe_value = int(aussenhoehe)  # Ensure this is correctly converted to an integer
            lager_queryset = lager_queryset.filter(
                aussenhöhe__gte=aussenhoehe_value - 10,
                aussenhöhe__lte=aussenhoehe_value + 10
            )
        except ValueError:
            pass  # In case a non-integer is entered

    if marke:
        # Adjust this depending on the field name in the 'Marke' model
        lager_queryset = lager_queryset.filter(marke__name__icontains=marke)

    # Always order by '-dichtungstyp' after filtering
    lager = lager_queryset.order_by('-dichtungstyp')

    context = {
        'lager': lager,
    }

    return render(request, 'crm/crm-lagerbestand.html', context)

@staff_member_required
def crm_lager_erfassen(request):
    if request.method == 'POST':
        # Initialize the form with POST data
        lager_form = CRMLagerForm(request.POST or None)

        if lager_form.is_valid():
            # Save the CRMLager instance
            lager = lager_form.save()  # Save the CRMLager form and get the instance

            # Redirect to another view or display a success message after saving
            return redirect('store:crm_lagerbestand')  # Change to the appropriate URL for your project

    else:
        # Initialize an empty form for GET request
        lager_form = CRMLagerForm()

    return render(request, 'crm/crm-lager-erfassen.html', {
        'lager_form': lager_form,
    })

@staff_member_required
def crm_lagerbestand_bearbeiten(request, pk):
    # Fetch the specific CRMLager instance using the primary key (pk)
    lager = get_object_or_404(CRMLager, pk=pk)

    if request.method == 'POST':
        # Initialize the form with POST data and the existing CRMLager instance
        lager_form = CRMLagerBestandForm(request.POST, instance=lager)

        if lager_form.is_valid():
            # Save only the 'lagerbestand' field
            lager_form.save()

            # Redirect to the lager overview page after successful update
            return redirect('store:crm_lagerbestand')  # Adjust URL to match your project

    else:
        # Prepopulate the form with the existing CRMLager 'lagerbestand' data
        lager_form = CRMLagerBestandForm(instance=lager)

    return render(request, 'crm/crm-lager-erfassen.html', {
        'lager_form': lager_form,
    })



@staff_member_required
def crm_lager_bearbeiten(request, pk):
    # Fetch the specific CRMLager instance using the primary key (pk)
    lager = get_object_or_404(CRMLager, pk=pk)

    if request.method == 'POST':
        # Initialize the form with POST data and the existing CRMLager instance
        lager_form = CRMLagerForm(request.POST, instance=lager)

        if lager_form.is_valid():
            # Save the updated CRMLager instance
            lager = lager_form.save()

            # Redirect to a lager overview page after successful update
            return redirect('store:crm_lagerbestand')  # Change to the appropriate URL

    else:
        # Prepopulate the form with the existing CRMLager instance data
        lager_form = CRMLagerForm(instance=lager)

    return render(request, 'crm/crm-lager-erfassen.html', {
        'lager_form': lager_form,
    })


@staff_member_required
def crm_lager_löschen(request, pk):
    eintrag = get_object_or_404(CRMLager, pk=pk)
    eintrag.delete()
    messages.info(request, "Der Artikel des Lagers wurde gelöscht.")
    return redirect("store:crm_lagerbestand")   

# CRM Kunden ---------------------------------------------------------------

def crm_new_kunden(request):
    search_query = request.GET.get('search', '')
    crm_kanton = request.GET.get('crm_kanton', '')

    # Initialize the CRMAddressForm with GET data to display the dropdown
    address_form = CRMKundenForm(request.GET or None)

    kunden = Kunde.objects.all().order_by('-id')  # Start with all customers

    if search_query:
        kunden = kunden.filter(
            Q(firmenname__icontains=search_query) |
            Q(interne_nummer__icontains=search_query) |
            Q(kunde_address__crm_ort__icontains=search_query) |
            Q(kunde_address__crm_strasse__icontains=search_query) |
            Q(kunde_address__crm_kanton__icontains=search_query)
        )

    if crm_kanton:
        kunden = kunden.filter(
            kunde_address__crm_kanton__icontains=crm_kanton
        )

    # Verhindert doppelte Kundeneinträge bei mehreren passenden Adressen
    kunden = kunden.distinct().order_by('-id')

    context = {
        'kunden': kunden,
        'address_form': address_form,
    }

    return render(request, 'crm/crm-kunden.html', context)



@staff_member_required
def crm_new_kunde_erfassen(request):
    if request.method == 'POST':
        # Initialize forms with POST data
        kunde_form = CRMKundeForm(request.POST)
        address_form = CRMAddressForm(request.POST)
        kunde_form2 = CRMKundeRestForm(request.POST)

        if kunde_form.is_valid() and address_form.is_valid():
            # First, save the Kunde instance
            kunde = kunde_form.save()  # Save the Kunde form and get the instance

            # Save kunde_form2 with the updated kunde instance
            kunde_form2.instance = kunde
            kunde_form2.save()
            
            # Now assign the Kunde instance to the CRMAddress form before saving
            address = address_form.save(commit=False)  # Don't save yet
            address.kunde = kunde  # Assign the saved Kunde instance
            address.address_type = 'R'  # Set the address type (or whatever default)
            address.save()  # Now save the CRMAddress instance

            return redirect('store:crm_new_kunden')

    else:
        # Empty forms for GET request
        kunde_form = CRMKundeForm()
        address_form = CRMAddressForm()
        kunde_form2 = CRMKundeRestForm()

    return render(request, 'crm/crm-kunde-erfassen.html', {
        'kunde_form': kunde_form,
        'address_form': address_form,
        'kunde_form2' : kunde_form2,
    })



@staff_member_required
def crm_new_kunde_bearbeiten(request, pk):
    kunde = get_object_or_404(Kunde, pk=pk)
    try:
        address = CRMAddress.objects.get(kunde=kunde)  # Get the related address
    except CRMAddress.DoesNotExist:
        address = None  # Handle the case where the address does not exist

    if request.method == 'POST':
        # Initialize the forms with POST data and existing instances
        kunde_form = CRMKundeForm(request.POST, instance=kunde)
        kunde_form2 = CRMKundeRestForm(request.POST, instance=kunde)
        address_form = CRMAddressForm(request.POST, instance=address)

        if kunde_form.is_valid() and kunde_form2.is_valid() and address_form.is_valid():
            # Save kunde_form first
            kunde = kunde_form.save()

            # Save kunde_form2 with the updated kunde instance
            kunde_form2.instance = kunde
            kunde_form2.save()

            # Save the updated address
            address = address_form.save(commit=False)
            address.kunde = kunde  # Ensure the address is associated with the correct Kunde
            address.address_type = 'R'  # Keep the same address type or update as needed
            address.save()

            return redirect('store:crm_new_kunden')

    else:
        # Prepopulate the forms with the existing Kunde and Address data
        kunde_form = CRMKundeForm(instance=kunde)
        kunde_form2 = CRMKundeRestForm(instance=kunde)
        address_form = CRMAddressForm(instance=address)

    return render(request, 'crm/crm-kunde-bearbeiten.html', {
        'kunde_form': kunde_form,
        'kunde_form2': kunde_form2,
        'address_form': address_form,
    })


@staff_member_required
def cms_crm_adresse_bearbeiten(request, pk):
    # Fetch the Kunde object (the customer)
    kunde = get_object_or_404(Kunde, pk=pk)

    # Try to get the existing address or create a new one for the Kunde
    try:
        address = CRMAddress.objects.get(kunde=kunde)
    except CRMAddress.DoesNotExist:
        address = CRMAddress(kunde=kunde)  # Create a new address with the linked Kunde

    if request.method == "POST":
        form = CRMAddressForm(request.POST, instance=address)
        if form.is_valid():
            # Manually set the Kunde field and address_type to 'R' before saving
            address = form.save(commit=False)
            address.kunde = kunde  # Set the 'kunde' field
            address.address_type = 'R'  # Set the default address_type to 'R'
            address.save()
            messages.success(request, "Address updated successfully.")
            return redirect('store:crm_new_kunden')
        else:
            messages.error(request, "Error updating address.")
    else:
        form = CRMAddressForm(instance=address)

    context = {
        'form': form,
        'kunde': kunde,
    }
    return render(request, 'crm/crm-adresse-bearbeiten.html', context)

@staff_member_required
def cms_crm_kunde_löschen(request, pk):
    eintrag = get_object_or_404(Kunde, pk=pk)
    eintrag.delete()
    messages.info(request, "Der Kunde wurde gelöscht.")
    return redirect("store:crm_new_kunden") 

@staff_member_required
def crm_update_last_service(request, pk):
    # Fetch the specific customer by primary key
    kunde = get_object_or_404(Kunde, pk=pk)  

    if request.method == 'POST':
        # Pass POST data and instance of the Kunde to update
        form = CRMLastService(request.POST, instance=kunde)

        if form.is_valid():
            form.save()  # Save the last_service update
            messages.success(request, "Letzter Service erfolgreich aktualisiert!")
            return redirect('store:crm_new_kunden')  # Redirect after successful save
        else:
            # Handle form validation errors
            messages.error(request, "Fehler beim Aktualisieren des letzten Service.")
    else:
        # Initialize the form with existing data for GET request
        form = CRMLastService(instance=kunde)

    return render(request, 'crm/crm-update-last-service.html', {
        'form': form,
        'kunde': kunde,
    })


# CRM END

# ELEMENTE

@staff_member_required
def cms_elemente(request, pk):
    kunde_data = Kunde.objects.get(pk=pk)
    search_query = request.GET.get('search', '')

    if search_query:
        elemente = Elemente.objects.filter(
            Q(dichtungen__titel__icontains=search_query) |
            Q(kuehlposition__icontains=search_query) |
            Q(aussenbreite__icontains=search_query) |
            Q(aussenhöhe__icontains=search_query) |
            Q(elementnr__icontains=search_query) |
            Q(artikel__name__icontains=search_query) |  # Added search for artikel name
            Q(artikel__artikelnr__icontains=search_query)  # Added search for artikel artikelnr
        ).filter(kunde=kunde_data).order_by('kunde', 'elementnr')
    else:
        elemente = Elemente.objects.filter(kunde=kunde_data).order_by('elementnr')

    context = {
        'elemente': elemente,
        'kunde_id': pk,
        'kunde_data': kunde_data,
    }
    return render(request, 'crm/cms-elemente.html', context)




def cms_elemente_duplicate(request, pk, elemente_pk):
    kunde = get_object_or_404(Kunde, pk=pk)
    
    # Fetch the Elemente instance to duplicate
    elemente_to_duplicate = get_object_or_404(Elemente, pk=elemente_pk, kunde=kunde)
    
    # Duplicate the specified Elemente instance
    elemente_to_duplicate.pk = None  # Reset the primary key to create a new object
    elemente_to_duplicate.elementnr = (elemente_to_duplicate.elementnr or 0) + 1  # Increment the elementnr
    elemente_to_duplicate.produkt = "DUPLIKAT"  # Update produkt field
    elemente_to_duplicate.kuehlposition = elemente_to_duplicate.kuehlposition
    elemente_to_duplicate.save()

    # Link the duplicated Elemente instance to the Kunde
    elemente_to_duplicate.kunde.add(kunde)

    # Redirect to the Elemente list
    messages.success(request, f"Das Element wurde erfolgreich dupliziert mit der neuen Elemente-Nr. {elemente_to_duplicate.elementnr}.")
    return redirect('store:cms_elemente', pk=kunde.pk)




@staff_member_required
def cms_elemente_create(request, pk):
    print("Bezeichnungen:", Bezeichnung.objects.all())

    kunde = get_object_or_404(Kunde, pk=pk)  # Get the Kunde instance
    artikel_liste = Artikel.objects.all()  # Get all Artikel
    form = ElementeCreateForm(request.POST or None)

    if request.method == "POST":
        artikel_name = request.POST.get('artikel_name', None)  # Get selected Artikel name
        selected_artikel = None

        # Attempt to find the Artikel by name (or another identifier, e.g., artikelnr)
        if artikel_name:
            try:
                selected_artikel = Artikel.objects.get(artikelnr=artikel_name)  # Adjust field if needed
            except Artikel.DoesNotExist:
                messages.error(request, f"Der Artikel '{artikel_name}' wurde nicht gefunden.")

        if form.is_valid():
            # Create the Elemente instance but do not save to the database yet
            elemente_instance = form.save(commit=False)

            # Link the selected Artikel if found
            if selected_artikel:
                elemente_instance.artikel = selected_artikel
                elemente_instance.aussenbreite = selected_artikel.aussenbreite
                elemente_instance.aussenhöhe = selected_artikel.aussenhöhe
            else:
                messages.error(request, "Bitte wählen Sie einen gültigen Artikel aus.")

            # Save the Elemente instance
            elemente_instance.save()

            # Link the Elemente instance to the Kunde
            elemente_instance.kunde.add(kunde)

            messages.success(request, "Das Element wurde erfolgreich erstellt.")
            return redirect('store:cms_elemente', pk=pk)
        else:
            messages.error(request, "Ein Fehler ist aufgetreten. Bitte überprüfen Sie die Eingaben.")

    context = {
        'form': form,
        'artikel_liste': artikel_liste,
        'kunde_id': pk,
    }
    return render(request, 'crm/cms-elemente-erfassen.html', context)

@staff_member_required
def fetch_artikel(request):
    query = request.GET.get('query', '').strip()

    if query:
        artikel_list = Artikel.objects.filter(artikelnr__icontains=query)[:5]  # Get first 5 matching results
    else:
        artikel_list = []

    artikel_data = [{'artikelnr': artikel.artikelnr, 'name': artikel.name} for artikel in artikel_list]
    return JsonResponse(artikel_data, safe=False)


@staff_member_required
def cms_elemente_edit(request, pk, cpk):
    element = get_object_or_404(Elemente, pk=pk)  # Get the Elemente instance
    artikel_liste = Artikel.objects.all()  # Get all Artikel
    form = ElementeCreateForm(request.POST or None, instance=element)

    if request.method == "POST":
        artikel_name = request.POST.get('artikel_name', None)  # Get selected Artikel name
        selected_artikel = None

        # Attempt to find the Artikel by artikelnr or another identifier
        if artikel_name:
            try:
                selected_artikel = Artikel.objects.get(artikelnr=artikel_name)  # Adjust field if needed
            except Artikel.DoesNotExist:
                messages.error(request, f"Der Artikel '{artikel_name}' wurde nicht gefunden.")

        if form.is_valid():
            # Update the Elemente instance but do not save to the database yet
            element = form.save(commit=False)

            # Link the selected Artikel if found
            if selected_artikel:
                element.artikel = selected_artikel
                element.aussenbreite = selected_artikel.aussenbreite
                element.aussenhöhe = selected_artikel.aussenhöhe
            else:
                messages.error(request, "Bitte wählen Sie einen gültigen Artikel aus.")

            # Save the Elemente instance
            element.save()

            messages.success(request, "Das Element wurde erfolgreich aktualisiert.")
            return redirect('store:cms_elemente', pk=cpk)
        else:
            messages.error(request, "Ein Fehler ist aufgetreten. Bitte überprüfen Sie die Eingaben.")

    context = {
        'form': form,
        'artikel_liste': artikel_liste,
        'element': element,
        'cpk': cpk,
    }
    return render(request, 'crm/cms-elemente-bearbeiten.html', context)



@staff_member_required
def cms_elemente_löschen(request, pk, cpk):
    eintrag = get_object_or_404(Elemente, pk=pk)
    eintrag.delete()
    messages.info(request, "Der Eintrag wurde gelöscht.")
    return redirect('store:cms_elemente', pk=cpk)   
#ELEMENTE END

@staff_member_required
def generate_lieferschein(request, bestellung_id):
    try:
        bestellung = Elemente_Bestellungen.objects.get(id=bestellung_id)
    except Elemente_Bestellungen.DoesNotExist:
        messages.error(request, "Bestellung not found")
        return redirect('store:elemente_bestellungen')

    # Path to the .docx template file
    template_path = os.path.join(settings.BASE_DIR, 'static/docx/template.docx')

    # Load the .docx template
    document = Document(template_path)

    # Replace `{{date}}` placeholder with today's date
    for paragraph in document.paragraphs:
        if "{{date}}" in paragraph.text:  # Search for `{{date}}`
            paragraph.text = paragraph.text.replace("{{date}}", date.today().strftime("%d.%m.%Y"))

    # Handle the `{{lieferschein.table}}` placeholder
    for paragraph in document.paragraphs:
        if "{{lieferschein.table}}" in paragraph.text:  # Search for `{{lieferschein.table}}`
            paragraph.text = ""  # Clear placeholder text

            # Create a table with the required structure
            table = document.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            # Add header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Kunden Nr."
            hdr_cells[1].text = "Betrieb/Person"
            hdr_cells[2].text = "Elemente Nr."
            hdr_cells[3].text = "Montage"
            hdr_cells[4].text = "Bemerkung"

            # Add data row
            row_cells = table.add_row().cells
            row_cells[0].text = bestellung.kunden_nr or "N/A"
            row_cells[1].text = f"{bestellung.betrieb_person}\n{bestellung.adresse}\n{bestellung.plz} {bestellung.ort}"
            row_cells[2].text = bestellung.elemente_nr or "N/A"
            row_cells[3].text = "Ja" if bestellung.montage == "Ja" else "Nein"
            row_cells[4].text = bestellung.bemerkung or "Keine Bemerkung"

            # Insert the table directly after the placeholder paragraph
            paragraph._element.addnext(table._tbl)
            break

    # Ensure the temp directory exists
    temp_dir = os.path.join(settings.BASE_DIR, 'temp')
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)

    # Save the generated Word document
    output_docx_path = os.path.join(temp_dir, f'Lieferschein_{bestellung.kunden_nr}.docx')
    document.save(output_docx_path)

    # Serve the Word document as a response
    with open(output_docx_path, 'rb') as docx_file:
        response = HttpResponse(docx_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=Lieferschein_{bestellung.kunden_nr}.docx'

    # Clean up temporary file
    os.remove(output_docx_path)

    # Add a success message
    messages.success(request, "Lieferschein wurde heruntergeladen.")
    return response





@staff_member_required
def download_qr_code(request):
    url = request.build_absolute_uri(reverse('store:bestellformular_cart'))

    qr = qrcode.QRCode(
        version=1,  # Controls the size of the QR Code
        error_correction=qrcode.constants.ERROR_CORRECT_L,  # Less error correction
        box_size=10,  # Size of each box in pixels
        border=4,  # Border size
    )
    
    # Add the URL data to the QR code
    qr.add_data(url)
    qr.make(fit=True)

    # Create an image from the QR code
    img = qr.make_image(fill='black', back_color='white')

    # Save the image to a BytesIO object
    buffer = BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)

    # Return the image as an HTTP response with content-disposition to download
    response = HttpResponse(buffer, content_type="image/png")
    response['Content-Disposition'] = 'attachment; filename="qr_code_bestellformular.png"'
    return response

@staff_member_required
def qr_code_view(request):
    url = request.build_absolute_uri(reverse('store:bestellformular_cart'))

    # Create the QR code instance
    qr = qrcode.QRCode(
        version=1,  # Controls the size of the QR Code
        error_correction=qrcode.constants.ERROR_CORRECT_L,  # Less error correction
        box_size=10,  # Size of each box in pixels
        border=4,  # Border size
    )
    
    # Add the URL data to the QR code
    qr.add_data(url)
    qr.make(fit=True)

    # Create an image from the QR code
    img = qr.make_image(fill='black', back_color='white')

    # Save the image to a BytesIO object
    buffer = BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)

    # Convert the image to base64 to embed in HTML
    image_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')

    # Pass the base64 image to the template
    return render(request, 'crm/qr-code.html', {'qr_code_base64': image_base64})


@staff_member_required
def delete_kunde_user_relationship(request, pk):
    # Get the Kunde object by primary key
    kunde = get_object_or_404(Kunde, pk=pk)
    
    # Unlink the user
    kunde.user = None
    kunde.save()
    
    # Redirect back to the customer list or another page
    return redirect(reverse('store:cms_kunden'))





def danke(request):
    context = {
    }
    return render (request, 'crm/danke.html', context)


#CMS



@staff_member_required
def cms_elemente_statistik(request):
    # Initialize the queryset
    elemente = Elemente.objects.all()
    
    # Capture individual search criteria
    produkt = request.GET.get('produkt', '')
    kuehlposition = request.GET.get('kuehlposition', '')
    aussenbreite = request.GET.get('aussenbreite', '')
    aussenhöhe = request.GET.get('aussenhöhe', '')
    serie = request.GET.get('serie', '')
    modell = request.GET.get('modell', '')
    typ = request.GET.get('typ', '')
    name = request.GET.get('name', '')
    
    # Capture sort_by parameter, default to sorting by 'kunde__firmenname'
    sort_by = request.GET.get('sort', 'kunde')

    # Filter the queryset based on search criteria
    if produkt:
        elemente = elemente.filter(produkt__icontains=produkt)
        if not elemente.exists():
            elemente = Elemente.objects.filter(dichtungen__titel__icontains=produkt)

    if kuehlposition:
        elemente = elemente.filter(kuehlposition__icontains=kuehlposition)
    
    if aussenbreite:
        try:
            aussenbreite_value = int(aussenbreite)
            elemente = elemente.filter(aussenbreite=aussenbreite_value)
        except ValueError:
            pass  # Ignore invalid integer input
    
    if aussenhöhe:
        try:
            aussenhöhe_value = int(aussenhöhe)
            elemente = elemente.filter(aussenhöhe=aussenhöhe_value)
        except ValueError:
            pass  # Ignore invalid integer input
    
    if serie:
        elemente = elemente.filter(elemente_objekte__serie__icontains=serie)
    
    if modell:
        elemente = elemente.filter(elemente_objekte__modell__icontains=modell)
    
    if typ:
        elemente = elemente.filter(elemente_objekte__typ__icontains=typ)
    
    if name:
        elemente = elemente.filter(elemente_objekte__name__icontains=name)

    # Apply sorting based on the 'sort' parameter
    if sort_by == 'kunde':
        elemente = elemente.order_by('kunde__firmenname')
    elif sort_by == 'elementnr':
        elemente = elemente.order_by('elementnr')
    else:
        elemente = elemente.order_by('kunde__firmenname')

    # Use distinct() to avoid duplicates if multiple criteria match
    elemente = elemente.distinct()

    # Calculate total laufmeter with fallback for missing dimensions
    total_laufmeter = 0
    for element in elemente:
        breite = element.aussenbreite or (element.artikel.aussenbreite if element.artikel else None)
        hoehe = element.aussenhöhe or (element.artikel.aussenhöhe if element.artikel else None)
        if breite and hoehe:
            total_laufmeter += (2 * (breite + hoehe)) / 1000  # Convert to meters

    # Count the number of Elemente records
    elemente_count = elemente.count()

    # Pass total laufmeter and filtered elements to the context
    context = {
        'elemente': elemente,
        'total_laufmeter': total_laufmeter,  # Pass total laufmeter to template
        'elemente_count': elemente_count,  # Pass the count of elements to template
        'sort_by': sort_by,  # Pass sorting parameter to the template
    }

    return render(request, 'crm/cms-elemente-statistik.html', context)




# for Bezeichnung section
def bezeichnung_list(request):
    search_query = request.GET.get('search', '')
    if search_query:
        bezeichnungen = Bezeichnung.objects.filter(
            Q(name__icontains=search_query)
        ).order_by('-id')
    else:
        bezeichnungen = Bezeichnung.objects.all().order_by('name')
    return render(request, 'crm/bezeichnungen.html', {'bezeichnungen': bezeichnungen})


def bezeichnung_create(request):
    if request.method == 'POST':
        form = BezeichnungForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('store:bezeichnung_list')  # Redirect to the list of Bezeichnung objects after creation
    else:
        form = BezeichnungForm()

    return render(request, 'crm/bezeichnung-create.html', {'form': form})

def bezeichnung_edit(request, pk):
    bezeichnung = get_object_or_404(Bezeichnung, pk=pk)
    if request.method == 'POST':
        form = BezeichnungForm(request.POST, instance=bezeichnung)
        if form.is_valid():
            form.save()
            return redirect('store:bezeichnung_list')
    else:
        form = BezeichnungForm(instance=bezeichnung)
    return render(request, 'crm/bezeichnung-edit.html', {'form': form})

def bezeichnung_delete(request, pk):
    bezeichnung = get_object_or_404(Bezeichnung, pk=pk)
    bezeichnung.delete()
    messages.info(request, "Die Bezeichnung wurde gelöscht.")
    return redirect('store:bezeichnung_list')


@staff_member_required
def cms_elemente_objekte_löschen(request, pk, epk, cpk):
    eintrag = get_object_or_404(Objekte, pk=pk)
    eintrag.delete()
    messages.info(request, "Das Objekt wurde gelöscht.")
    return redirect("store:cms_elemente_objekte", pk=epk, cpk=cpk)


@staff_member_required
def cms_elemente_objekte_bearbeiten(request, pk, epk, cpk):
    # Retrieve the Elemente instance (as before)
    elemente_instance = get_object_or_404(Elemente, pk=epk)

    # Retrieve the existing Objekte instance that we want to edit
    objekte_instance = get_object_or_404(Objekte, pk=pk)

    # Initialize the form with the existing objekte_instance
    form = ElementeObjekteCreateForm(request.POST or None, instance=objekte_instance)

    if request.method == "POST":
        if form.is_valid():
            # Save the form and update the objekte_instance
            objekte_instance = form.save(commit=False)
            objekte_instance.save()

            # Ensure the relationship between objekte_instance and elemente_instance is maintained
            objekte_instance.objekte.set([elemente_instance])  # Maintain the association
            objekte_instance.save()

            messages.success(request, "Objekt successfully updated and associated with Elemente.")
            return redirect("store:cms_elemente_objekte", pk=epk, cpk=cpk)
        else:
            messages.error(request, "Error in form submission. Please check the fields and try again.")

    context = {
        'form': form,
        'pk': pk,
        'epk': epk,
        'cpk': cpk,
    }
    return render(request, 'crm/cms-elemente-objekte-erfassen.html', context)


@staff_member_required
def cms_elemente_objekte_erfassen(request, pk, cpk):
    # Ensure we get a single instance of Elemente
    elemente_instance = get_object_or_404(Elemente, pk=pk)  # Use get_object_or_404 for a single object

    form = ElementeObjekteCreateForm(request.POST or None)
    
    if request.method == "POST":
        if form.is_valid():
            objekte_instance = form.save(commit=False)
            objekte_instance.save()
            # Associate the newly created objekte with the given elemente
            objekte_instance.objekte.set([elemente_instance])  # Pass a list to set
            objekte_instance.save()
            messages.success(request, "Objekt successfully created and associated with Elemente.")
            return redirect("store:cms_elemente_objekte", pk=pk, cpk=cpk)
        else:
            messages.error(request, "Error in form submission. Please check the fields and try again.")
    
    context = {
        'form': form,
        'pk': pk,
        'cpk': cpk,
    }
    return render(request, 'crm/cms-elemente-objekte-erfassen.html', context)

@staff_member_required
def cms_elemente_objekte(request, pk, cpk):
    kunde = Kunde.objects.get(pk=cpk)
    elemente_data = get_object_or_404(Elemente, pk=pk)
    objekte = Objekte.objects.filter(objekte=elemente_data).order_by('id')

    context = {
        'kunde': kunde,
        'cpk':cpk,
        'pk': pk,
        'objekte': objekte,
    }
    return render(request, 'crm/cms-elemente-objekte.html', context)


def anleitung_videos(request):
    context = {
    }
    return render (request, 'anleitung-videos.html', context)




def error_404(request, exception):
        data = {}
        return render(request,'404.html', data)

#Hauptseite
def home(request):
    search_query = request.GET.get('search', '')
    if search_query:
        items = Item.objects.filter(titel__icontains=search_query)
    else:
        items = Item.objects.all()
        if request.method =="POST":
            vorname = request.POST['message-vorname']
            nachname = request.POST['message-name']
            firma = request.POST['message-firma']
            email = request.POST['message-email']
            nachricht = request.POST['message']
            telefon = request.POST['message-phone']
            anrede = request.POST['message-anrede']

            subject = 'gastrodichtung.ch Nachricht von ' + ' '+ firma + ' ' + vorname + ' ' + nachname
            template = render_to_string('shop/kontakt-email.html', {
                'anrede' : anrede,
                'vorname': vorname, 
                'nachname': nachname,
                'firma': firma,
                'email': email,
                'telefon': telefon,
                'nachricht': nachricht,         
                 })
            
            #send email for order
            email = EmailMessage(
                subject,
                template,
                email,
                ['bestellungen@gastrodichtung.ch', 'livio.bonetta@geboshop.ch'],
            )

            email.fail_silently=False
            email.content_subtype = "html"
            email.send()



            context = {
                'items': items,
                'vorname' : vorname,
            }
            return render(request, 'dichtungen.html', context)
        else:
            marken = Marke.objects.all()
            context = {
                'items': items,
                'marken' : marken, 
            }
            return render(request, 'dichtungen.html', context)

    return render(request, 'dichtungen.html', context)

def searchbar(request):
    search_query = request.GET.get('search', '')
    if search_query:
        items = Item.objects.filter(Q(titel__icontains=search_query) | Q(artikelnr__icontains=search_query) | Q(kategorie__name__icontains=search_query) | Q(subkategorie__sub_name__icontains=search_query))
    else:
        items = Item.objects.all()


    context = {'items': items}
    return render(request, 'shop/searchbar.html', context)


def impressum(request):
    context = {}
    return render(request, 'impressum.html', context)


#Kontaktseite
def kontakt(request):

    if request.method =="POST":
        vorname = request.POST['message-surname']
        nachname = request.POST['message-name']
        firma = request.POST['message-company']
        email = request.POST['message-email']
        nachricht = request.POST['message']
        telefon = request.POST['message-phone']
        anrede = request.POST['message-anrede']

        subject = 'Nachricht von ' + ' '+ firma + ' ' + vorname + ' ' + nachname
        template = render_to_string('shop/kontakt-email.html', {
            'anrede' : anrede,
            'vorname': vorname, 
            'nachname': nachname,
            'firma': firma,
            'email': email,
            'telefon': telefon,
            'nachricht': nachricht,         
             })
        
        #send email for order
        email = EmailMessage(
            subject,
            template,
            email,
            email_master,
        )

        email.fail_silently=False
        email.content_subtype = "html"
        email.send()

        context = {
            'message_kontakt' : 'Die Nachricht wurde erfolgreich gesendet.'
             }
        return render(request, 'kontakt.html', context)
    
    else:
        context = { }
        return render(request, 'kontakt.html', context)

#Kontaktseite END



def firma(request):
    context = { }
    return render(request, 'montage.html', context)



#Marke Übersicht
def marke(request):
    marken = Marke.objects.all()
    context = { 
        'marken' : marken, 
        }
    return render(request, 'shop/store-marke.html', context)


#Marke Store
def HomeMarkeView(request, cat_marke):
    seo_marke = get_object_or_404(Marke, slug=cat_marke)
    marke_products = Item.objects.filter(item_marken__slug=cat_marke)
    context = { 
        'marke_products' : marke_products, 
        'cat_marke': cat_marke.replace('-', ''),
        'seo_marke' : seo_marke,
        }
    return render(request, 'shop/store-marke-details.html', context)

#Produktübersichten
def HomeProduktView(request, cats):
    category_products = Item.objects.filter(kategorie__slug=cats)
    seo_cat = get_object_or_404(Category, slug=cats)
    context = { 
        'category_products' : category_products, 
        'cats': cats.replace('-', ''),
        'seo_cat' : seo_cat,
        }
    return render(request, 'shop/store.html', context)


#Produktübersichten
def HomeSubProduktView(request, subcats, cats):
    category_products = Item.objects.filter(kategorie__slug=cats)
    subcategory_products = Item.objects.filter(subkategorie__sub_slug=subcats)
    seo_subcat = get_object_or_404(Subcategory, sub_slug=subcats)
    context = { 
        'category_products' : category_products, 
        'cats': cats.replace('-', ''),
        'subcategory_products' : subcategory_products, 
        'subcats': subcats.replace('-', ''),
        'seo_subcat' : seo_subcat,
        }
    return render(request, 'shop/store-subcat.html', context)


def product_detail(request, slug):
    item = get_object_or_404(Item, slug=slug)
    form = AussenmassForm(request.POST or None)

    if request.method == "POST" and "email" in request.POST:
        # Allgemeine Produktinfos
        produkt = item.titel
        artikelnr = item.artikelnr
        produktlink = request.build_absolute_uri()

        # Felder aus dem Formular
        laufmeter = request.POST.get('laufmeter')
        aussenbreite = request.POST.get('aussenbreite')
        aussenhöhe = request.POST.get('aussenhöhe')
        anzahl = request.POST.get('anzahl')
        name = request.POST.get('name')
        email = request.POST.get('email')
        telefon = request.POST.get('telefon', '')
        nachricht = request.POST.get('nachricht', '')

        # E-Mail-Betreff
        subject = f"Produktanfrage zu {item.kategorie} {artikelnr}"

        # Template-Daten vorbereiten
        template_data = {
            'produkt_kategorie': item.kategorie,
            'produkt': produkt,
            'artikelnr': artikelnr,
            'produktlink': produktlink,
            'aussenbreite': aussenbreite,
            'aussenhöhe': aussenhöhe,
            'laufmeter': laufmeter,
            'anzahl': anzahl,
            'name': name,
            'email': email,
            'telefon': telefon,
            'nachricht': nachricht,
        }

        # HTML-E-Mail rendern
        template = render_to_string('emails/produktanfrage-email.html', template_data)

        email_message = EmailMessage(
            subject,
            template,
            settings.EMAIL_HOST_USER,
            email_master,
        )
        email_message.fail_silently = False
        email_message.content_subtype = "html"
        email_message.send()

        return redirect('store:anfrage_danke')  # oder was auch immer dein URL-Name ist

    # Normale GET-Anzeige
    context = {
        'item': item,
        'form': form,
    }
    return render(request, 'shop/descriptions.html', context)


def anfrage_danke_view(request):
    return render(request, 'shop/anfrage-danke.html')

    # Normale GET-Anzeige
    context = {
        'item': item,
        'form': form,
    }
    return render(request, 'shop/descriptions.html', context)

def anfrage_danke_view(request):
    return render(request, 'shop/anfrage-danke.html')

#Produktbeschreibungen
def weitere_product_detail(request, slug):
    item = get_object_or_404(Item, slug=slug)
    form = AussenmassForm(request.POST or None)

    if request.method == "POST" and "email" in request.POST:
        # Allgemeine Produktinfos
        produkt = item.titel
        artikelnr = item.artikelnr
        produktlink = request.build_absolute_uri()

        # Formularfelder
        laufmeter = request.POST.get('laufmeter')
        aussenbreite = request.POST.get('aussenbreite')
        aussenhöhe = request.POST.get('aussenhöhe')
        anzahl = request.POST.get('anzahl')
        name = request.POST.get('name')
        email = request.POST.get('email')
        telefon = request.POST.get('telefon', '')
        nachricht = request.POST.get('nachricht', '')

        # E-Mail-Betreff
        subject = f"Produktanfrage zu {item.kategorie} {artikelnr}"

        # Template-Daten
        template_data = {
    'produkt_kategorie': item.kategorie,
    'produkt_subkategorie': getattr(item, 'subkategorie', None),
    'produkt': produkt,
    'artikelnr': artikelnr,
    'produktlink': produktlink,
    'aussenbreite': aussenbreite,
    'aussenhöhe': aussenhöhe,
    'laufmeter': laufmeter,
    'anzahl': anzahl,
    'name': name,
    'email': email,
    'telefon': telefon,
    'nachricht': nachricht,
}


        # E-Mail versenden
        template = render_to_string('emails/produktanfrage-email.html', template_data)
        email_message = EmailMessage(
            subject,
            template,
            settings.EMAIL_HOST_USER,
            email_master,
        )
        email_message.fail_silently = False
        email_message.content_subtype = "html"
        email_message.send()

        return redirect('store:anfrage_danke')  # URL-Name sicherstellen!

    context = {
        'item': item,
        'form': form,
    }
    return render(request, 'shop/descriptions_sub.html', context)


@login_required
def add_to_cart(request, slug, pk):
    
    #check if it's post or not
    if request.method == "POST":
        aussenbreite = request.POST['aussenbreite']
        aussenhöhe = request.POST['aussenhöhe']
        anzahl = int(request.POST['anzahl'])
        
        #get the right product
        item = get_object_or_404(Item, slug=slug)
        
        #check if the order exists
        order_qs = Order.objects.filter(user=request.user, ordered=False)
        if order_qs.exists():
            order = order_qs[0]

            #check if there is existing order item
            order_item, created = OrderItem.objects.get_or_create(
                item=item,
                user=request.user,
                ordered=False,
                aussenbreite=aussenbreite,
                aussenhöhe=aussenhöhe,
                )

            #check if the order item is in the order

            #if the order item exist adjust quantity
            if order.items.filter(item__slug=item.slug, aussenbreite=aussenbreite, aussenhöhe=aussenhöhe).exists():
                order_item.quantity += anzahl
                order_item.save()
                messages.info(request, "Die Menge wurde aktualisiert.")
                return redirect("store:order_summary")
            
            #if the order item not exist add it into order
            else:
                #these below steps are working, dont change
                order_item.aussenbreite = aussenbreite
                order_item.aussenhöhe = aussenhöhe
                order_item.quantity = anzahl
                order_item.save()
                order.items.add(order_item)
                messages.info(request, "Das Produkt wurde erfolgreich in den Warenkorb gelegt.")
                return redirect("store:order_summary")

        
        #the order not existing for magnetdichtung and pvcohne
        else:
            ordered_date = timezone.now()
            order = Order.objects.create(user=request.user, ordered_date=ordered_date)
            if request.method == "POST":
                aussenbreite = request.POST['aussenbreite']
                aussenhöhe = request.POST['aussenhöhe']
                anzahl = request.POST['anzahl']
                #these below steps are working, dont change
                order_item, created = OrderItem.objects.get_or_create(
                    item=item,
                    user=request.user,
                    ordered=False,
                    aussenbreite=aussenbreite,
                    aussenhöhe=aussenhöhe,
                    quantity=anzahl,
                        )
                order.items.add(order_item)
                messages.info(request, "Das Produkt wurde erfolgreich in den Warenkorb gelegt.")
                return redirect("store:order_summary")
            
            #the order not existing for gummidichtungen
            else:
                aussenbreite = 250
                aussenhöhe = 250
                ordered_date = timezone.now()
                order = Order.objects.create(user=request.user, ordered_date=ordered_date)
                #these below steps are working, dont change
                order_item, created = OrderItem.objects.get_or_create(
                    item=item,
                    user=request.user,
                    ordered=False,
                    aussenbreite=aussenbreite,
                    aussenhöhe=aussenhöhe,
                        )
                order.items.add(order_item)
                messages.info(request, "Das Produkt wurde erfolgreich in den Warenkorb gelegt.")
                return redirect("store:order_summary")

    
    #if it's not request.method == "POST"
    else:
        #get the right product
        item = get_object_or_404(Item, slug=slug)
        
        #check if the order exists
        order_qs = Order.objects.filter(user=request.user, ordered=False)
        if order_qs.exists():
            order = order_qs[0]

            #check if there is existing order item
            order_item, created = OrderItem.objects.get_or_create(
                item=item,
                user=request.user,
                ordered=False,
                pk=pk,
                )

            #check if the order item is in the order

            #if the order item exist adjust quantity
            if order.items.filter(item__slug=item.slug, pk=pk).exists():
                order_item.quantity +=1
                order_item.save()
                messages.info(request, "Die Menge wurde aktualisiert.")
                return redirect("store:order_summary")
            else:
                pass
                #this case not existing
        else:
            pass
            #this case not existing

#Kundendichtungen
@login_required
def mydichtungen(request):
    try:
        # Attempt to fetch the profile for the user
        kunde = Kunde.objects.get(user=request.user)
        allelements = Elemente.objects.filter(kunde=request.user.profile)
    except Kunde.DoesNotExist:
        # If no profile exists, set allelements to an empty list
        kunde = None
        allelements = []

    if request.method == "POST":
        mydslug = request.POST['mydslug']
        aussenbreite = request.POST['aussenbreite']
        aussenhöhe = request.POST['aussenhöhe']
        anzahl = request.POST['anzahl']
        element = request.POST['elementnr']
        item = get_object_or_404(Item, slug=mydslug)
        
        # Create or get the order item
        orderitem, created = OrderItem.objects.get_or_create(
            item=item,
            user=request.user,
            ordered=False,
            aussenbreite=aussenbreite, 
            aussenhöhe=aussenhöhe,
            element=element,
        )
        # Redirect to add_to_cart page
        return redirect(
            "store:add_to_cart_myd",
            slug=mydslug,
            pk=item.pk,
            aussenbreite=aussenbreite,
            aussenhöhe=aussenhöhe,
            anzahl=anzahl,
            element=element,
        )
    else:
        orderitem = None  # Initialize orderitem for context

    # Render the template with context
    context = { 
        'allelements': allelements,
        'orderitem': orderitem,
        'kunde': kunde,
    }
    return render(request, 'shop/mydichtungen.html', context)


#add to cart for mydichtungen
@login_required
def add_to_cart_myd(request, slug, pk, aussenbreite, aussenhöhe, anzahl, element):
    #get the right product
    item = get_object_or_404(Item, slug=slug)
        
    #check if the order exists
    order_qs = Order.objects.filter(user=request.user, ordered=False)
    if order_qs.exists():
        order = order_qs[0]

        #check if there is existing order item
        order_item, created = OrderItem.objects.get_or_create(
                item=item,
                user=request.user,
                ordered=False,
                aussenbreite=aussenbreite,
                aussenhöhe=aussenhöhe,
                element=element,
                )

            #check if the order item is in the order

        #if the order item exist adjust quantity
        if order.items.filter(item__slug=item.slug, aussenbreite=aussenbreite, aussenhöhe=aussenhöhe, element=element).exists():
            order_item.quantity += anzahl
            order_item.save()
            messages.info(request, "Die Menge wurde aktualisiert.")
            return redirect("store:order_summary")
            
        #if the order item not exist add it into order
        else:
            #these below steps are working, dont change
            order_item.aussenbreite = aussenbreite
            order_item.aussenhöhe = aussenhöhe
            order_item.quantity = anzahl
            order_item.element = element
            order_item.save()
            order.items.add(order_item)
            messages.info(request, "Das Produkt wurde erfolgreich in den Warenkorb gelegt.")
            return redirect("store:order_summary")

        
    #the order not existing for magnetdichtung and pvcohne
    else:
        ordered_date = timezone.now()
        order = Order.objects.create(user=request.user, ordered_date=ordered_date)
            
        #these below steps are working, dont change
        order_item, created = OrderItem.objects.get_or_create(
                    item=item,
                    user=request.user,
                    ordered=False,
                    aussenbreite=aussenbreite,
                    aussenhöhe=aussenhöhe,
                    element=element,
                        )
        order.items.add(order_item)
        messages.info(request, "Das Produkt wurde erfolgreich in den Warenkorb gelegt.")
        return redirect("store:order_summary")
        
@login_required
def remove_from_cart(request, slug, pk):
    item = get_object_or_404(Item, slug=slug)
    order_qs = Order.objects.filter(
        user=request.user, 
        ordered=False
    )
    if order_qs.exists():
        order = order_qs[0]
        #check if the order item is in the order
        if order.items.filter(item__slug=item.slug, pk=pk).exists():
            order_item = OrderItem.objects.filter(
                item=item,
                user=request.user,
                ordered=False,
                pk=pk,
            )[0]
            order.items.remove(order_item)
            order_item.delete()
            messages.info(request, "Das Produkt wurde aus dem Warenkorb gelöscht.")
            return redirect("store:order_summary")  
        else:
            messages.info(request, "Es hat kein Produkt im Warenkorb.")
            return redirect("store:product-detail", slug=slug)  
    else:
        messages.info(request, "Es gibt keine aktuelle Bestellung.")
        return redirect("store:product-detail", slug=slug)


@login_required
def remove_single_item_from_cart(request, slug, pk):
    item = get_object_or_404(Item, slug=slug)
    order_qs = Order.objects.filter(
        user=request.user, 
        ordered=False
    )
    if order_qs.exists():
        order = order_qs[0]
        #check if the order item is in the order
        if order.items.filter(item__slug=item.slug, pk=pk).exists():
            order_item = OrderItem.objects.filter(
                item=item,
                user=request.user,
                ordered=False,
                pk=pk,
            )[0]
            if order_item.quantity >1 :
                order_item.quantity -=1
                order_item.save()
            else: 
                order.items.remove(order_item)
                order_item.delete()
            messages.info(request, "Die Menge wurde angepasst.")
            return redirect("store:order_summary")  
        else:
            messages.info(request, "Es hat kein Produkt im Warenkorb.")
            return redirect("store:product-detail", slug=slug)  
    else:
        messages.info(request, "Es gibt keine aktuelle Bestellung.")
        return redirect("store:product-detail", slug=slug)

# Warenkorb
class OrderSummaryView(LoginRequiredMixin, View):
    def get(self, *args, **kwargs):
        try:
            order = Order.objects.get(user=self.request.user, ordered=False)
            total = order.get_total

            # Fetch shipping costs based on the total
            shipping1 = ShippingCost.objects.filter(price_from__lte=total, price_to__gt=total)

            if shipping1.exists():  # Check if shipping costs exist
                shipping_price = shipping1[0].shipping_price
                zuschlag = 18 if order.items.filter(item__subkategorie__sub_name="Duschdichtungen") else 0

                # Apply shipping cost and Zuschlag (if applicable)
                order.shippingcost = shipping_price + zuschlag
                order.save()

                context = {
                    'object': order,
                    'zuschlag': zuschlag if zuschlag > 0 else None,  # Only add if Zuschlag is applied
                }
            else:
                # Handle case where no shipping costs are found
                messages.warning(self.request, "Es wurden keine Versandkosten gefunden.")
                context = {
                    'object': order,
                    'zuschlag': None,
                    'total': total,
                }

            return render(self.request, 'shop/order_summary.html', context)

        except ObjectDoesNotExist:
            messages.info(self.request, "Es gibt keine Bestellung.")
            return redirect("store:home")

#Bestellung abschliessen
class FinalSummaryView(View):
    def get(self, *args, **kwargs):
        try:
            order = Order.objects.get(user=self.request.user, ordered=False)
            order_items = order.items.all()
            context = {
                'order_items' : order_items,
                'object' : order,
            }
            return render(self.request, 'shop/final_summary.html', context)
        except ObjectDoesNotExist:
            messages.info(self.request,"Es gibt keine Bestellung.")
            return redirect("store:home")   

    def post(self, *args, **kwargs):
        try:
            order = Order.objects.get(user=self.request.user, ordered=False)
            order_items = order.items.all()
        
            order_items.update(ordered=True)
            for item in order_items:
                item.save()

            order.discount = order.get_rabatt()
            order.discount_pct = order.user.profile.rabatt
            order.skonto = order.get_skonto()
            order.pre_total = order.get_total_pre_mwst()
            order.order_mwst = order.mwst()
            order.total = order.grandtotal()

            order.ordered = True
            order.payment = True
            order.save()

            subject = 'Auftragsbestätigung Bestellung 206505'+str(order.id) +' von Gastrodichtungen.ch'
            template = render_to_string('shop/bestellbestaetigung.html', {
                    'vorname': self.request.user.first_name, 
                    'nachname': self.request.user.last_name,
                    'datum' : str(order.ordered_date.strftime('%d' + '.' + ' %b' + ' %Y')),
                    'order_items' : order_items, 
                    'mwst' : order.mwst(),
                    'total': order.grandtotal(),
                    'object' : order,
                     })
                        
            #send email for order
            email = EmailMessage(
                    subject,
                    template,
                    settings.EMAIL_HOST_USER,
                    [self.request.user.email, 'livio.bonetta@geboshop.ch'],
                            )

            email.fail_silently=False
            email.content_subtype = "html"
            email.send()

            #redirect to success page
            return redirect('store:rechnung')

        except ObjectDoesNotExist:
            messages.error(request, "Es gibt keine Bestellung.")
            return redirect("store:home")



#Adresse und Lieferung
class CheckoutView(View):
    def get(self, *args, **kwargs):
        try:
            order = Order.objects.get(user=self.request.user, ordered=False)
            form = CheckoutForm()
            
            context = {
                            'form': form,
                            'order': order,
                        }

            billing_address_qs = Address.objects.filter(
                user=self.request.user,
                address_type='B',
            )
            if billing_address_qs.exists():
                context.update({'billing_address': billing_address_qs[0]})
            
            shipping_address_qs = ShippingAddress.objects.filter(
                user=self.request.user,
                address_type='S',
            )
            if shipping_address_qs.exists():
                context.update({'shipping_address': shipping_address_qs[0]})
            
            return render(self.request, 'shop/checkout.html', context)
        except ObjectDoesNotExist:
            messages.info(self.request, "Es gibt keine bestehende Bestellung.")
            return redirect("store:home")

    def post(self, *args, **kwargs):
        form = CheckoutForm(self.request.POST or None)
        try:
            user=self.request.user
            order = Order.objects.get(user=self.request.user, ordered=False)
            if form.is_valid():
                use_other_billing = form.cleaned_data.get(
                    'use_other_billing')
                different_shipping_address = form.cleaned_data.get(
                            'different_shipping_address')

                if use_other_billing:
                    rechnung_strasse = form.cleaned_data.get('rechnung_strasse')
                    rechnung_nr = form.cleaned_data.get('rechnung_nr')
                    rechnung_ort = form.cleaned_data.get('rechnung_ort')
                    rechnung_land = form.cleaned_data.get('rechnung_land')
                    rechnung_plz = form.cleaned_data.get('rechnung_plz')
                    rechnung_firmenname = form.cleaned_data.get('rechnung_firmenname')
                    rechnung_vorname = form.cleaned_data.get('rechnung_vorname')
                    rechnung_nachname = form.cleaned_data.get('rechnung_nachname')

                    billing_address = Address.objects.create(
                        user=self.request.user,
                        address_type='B',
                    )

                    billing_address.rechnung_strasse = rechnung_strasse
                    billing_address.rechnung_nr = rechnung_nr
                    billing_address.rechnung_ort = rechnung_ort
                    billing_address.rechnung_land = rechnung_land
                    billing_address.rechnung_plz = rechnung_plz
                    billing_address.address_type = 'B'
                    billing_address.save()
                    user.profile.firmenname = rechnung_firmenname
                    user.first_name = rechnung_vorname
                    user.last_name = rechnung_nachname
                    user.save()
                    user.profile.save()
                    order.billing_address = billing_address
                    order.save()

                    if different_shipping_address:
                            
                            lieferung_strasse = form.cleaned_data.get('lieferung_strasse')
                            lieferung_nr = form.cleaned_data.get('lieferung_nr')
                            lieferung_ort = form.cleaned_data.get('lieferung_ort')
                            lieferung_plz = form.cleaned_data.get('lieferung_plz')
                            lieferung_land = form.cleaned_data.get('lieferung_land')
                            lieferung_firmenname = form.cleaned_data.get('firmenname')
                            lieferung_vorname = form.cleaned_data.get('vorname')
                            lieferung_nachname = form.cleaned_data.get('nachname')

                            shipping_address = ShippingAddress.objects.create(
                                    user=self.request.user,
                                    address_type='S',

                                )

                            shipping_address.lieferung_strasse = lieferung_strasse
                            shipping_address.lieferung_nr = lieferung_nr
                            shipping_address.lieferung_ort = lieferung_ort                              
                            shipping_address.lieferung_plz = lieferung_plz
                            shipping_address.lieferung_land = lieferung_land
                            shipping_address.firmenname = lieferung_firmenname
                            shipping_address.vorname = lieferung_vorname
                            shipping_address.nachname = lieferung_nachname
                            shipping_address.address_type = 'S'
                            shipping_address.save()
                            order.shipping_address = shipping_address
                            order.save()
                            return redirect('store:payment')

                    else:
                        order.shipping_address = None
                        order.save()
                        return redirect('store:payment')

                else:
                    address_qs = Address.objects.filter(
                        user=self.request.user,
                        address_type='B',
                        )

                    if address_qs.exists():
                        billing_address = address_qs[0]
                        order.billing_address = billing_address
                        order.save()

                        if different_shipping_address:
                            lieferung_strasse = form.cleaned_data.get('lieferung_strasse')
                            lieferung_nr = form.cleaned_data.get('lieferung_nr')
                            lieferung_ort = form.cleaned_data.get('lieferung_ort')
                            lieferung_plz = form.cleaned_data.get('lieferung_plz')
                            lieferung_land = form.cleaned_data.get('lieferung_land')
                            lieferung_firmenname = form.cleaned_data.get('firmenname')
                            lieferung_vorname = form.cleaned_data.get('vorname')
                            lieferung_nachname = form.cleaned_data.get('nachname')
                            
                            shipping_address = ShippingAddress.objects.create(
                                    user=self.request.user,
                                    address_type='S',

                                )

                            shipping_address.lieferung_strasse = lieferung_strasse
                            shipping_address.lieferung_nr = lieferung_nr
                            shipping_address.lieferung_ort = lieferung_ort                              
                            shipping_address.lieferung_plz = lieferung_plz
                            shipping_address.lieferung_land = lieferung_land
                            shipping_address.firmenname = lieferung_firmenname
                            shipping_address.vorname = lieferung_vorname
                            shipping_address.nachname = lieferung_nachname
                            shipping_address.address_type = 'S'
                            shipping_address.save()
                            order.shipping_address = shipping_address
                            order.save()
                            return redirect('store:payment')

                        else:
                            order.shipping_address = None
                            order.save()
                            return redirect('store:payment')

                    else: 
                        return redirect('store:payment')                
            else: 
                pass
                        
        except ObjectDoesNotExist:
            messages.error(request, "Es gibt keine Bestellung.")
            return redirect("store:home")


#Adresse und Lieferung
class PaymentView(View):
    def get(self, *args, **kwargs):
        try:
            order = Order.objects.get(user=self.request.user, ordered=False)
            payment_form = PaymentForm()
            context = {
                'payment_form' : payment_form,
                'order': order,
            }
            return render(self.request, 'shop/zahlung.html', context)
        except ObjectDoesNotExist:
            messages.info(self.request, "Es gibt keine bestehende Bestellung.")
            return redirect("store:home")

    def post(self, *args, **kwargs):
        form = PaymentForm(self.request.POST or None)
        try:
            order = Order.objects.get(user=self.request.user, ordered=False)
            if form.is_valid():
                payment_option = form.cleaned_data.get('payment_option')
                order.payment_method = payment_option
                order.save()
                #success page redirect
                return redirect('store:final_summary')
            else:
                pass
                                
        except ObjectDoesNotExist:
            messages.error(request, "Es gibt keine Bestellung.")
            return redirect("store:home")

#Success Page
@login_required
def Rechnung(request):
    context = { }
    return render(request, 'shop/rechnung.html', context)

#Success Page2
@login_required
def email(request, pk):
    order = Order.objects.get(pk=pk)
    order_items = order.items.all()
    context = {
        'order_items' : order_items,
        'object' : order,
        }
    return render(request, 'shop/bestellung_uebersicht.html', context)

@login_required
def bestellungen(request):
    order = Order.objects.filter(user=request.user).order_by('-ordered_date')
    
    context = {
        'object' : order,
        }
    return render(request, 'shop/bestellungen.html', context)


@login_required
def einstellungen(request):
    address = Address.objects.get(user=request.user)
    user = request.user
    if request.method == "POST":
        firmenname = request.POST['firmenname']
        vorname = request.POST['vorname']
        nachname = request.POST['nachname']
        birthday = request.POST['birthday']
        phone = request.POST['phone']
        mobile = request.POST['mobile']
        user.profile.birthday = birthday
        user.profile.mobile = mobile
        user.profile.phone = phone
        user.profile.firmenname = firmenname
        user.profile.save()
        user.first_name = vorname
        user.last_name = nachname
        user.save()
        form = AddressForm(request.POST or None, instance=address)
        if form.is_valid():
            form.save()
            return redirect('store:einstellungen')

        else:
            messages.error(request, "Error")

    else: 
        form = AddressForm(request.POST or None, instance=address)

    context = {
        'form': form,
        'address' : address,
                }
    return render(request, 'shop/einstellungen.html', context)  

@login_required
def lieferadresse(request):
    address = ShippingAddress.objects.get(user=request.user)
    if request.method == "POST":
        form = ShippingAddressForm(request.POST or None, instance=address)
        if form.is_valid():
            form.save()
            return redirect('store:einstellungen')

        else:
            messages.error(request, "Error")

    else: 
        form = ShippingAddressForm(request.POST or None, instance=address)

    context = {
        'form': form,
        'address' : address,
                }
    return render(request, 'shop/lieferadresse.html', context)

@login_required
def create_lieferadresse(request):
    if request.method == "POST":
        form = ShippingAddressForm(request.POST or None)
        if form.is_valid():
            adresse = form.save(commit=False)
            adresse.user = request.user
            adresse.address_type = 'S'
            adresse.save()
            return redirect('store:einstellungen')

        else:
            messages.error(request, "Error")

    else: 
        form = ShippingAddressForm()

    context = {
        'form': form,
                }
    return render(request, 'shop/lieferadresse-erfassen.html', context)




#CMS related fields --- only for staff available
@staff_member_required
def cms(request):

    cat = Category.objects.first()

    context = {
    'cat' : cat
     }
    return render(request, 'cms.html', context)


@staff_member_required
def cms_bestellungen(request):
    order = Order.objects.filter(ordered=True).order_by('-ordered_date')
    context = {
            'order': order,
             }
    return render(request, 'cms-bestellungen.html', context)

#Success Page2
@staff_member_required
def cms_bestellung_confirmation(request, pk):
    order = Order.objects.get(pk=pk)
    order_items = order.items.all()
    context = {
        'order_items' : order_items,
        'object' : order,
        }
    return render(request, 'cms-bestellung-confirmation.html', context)




@staff_member_required
def cms_kennzahlen_webseite(request):
    import json
    import requests
    import ast
    
    api_request_today = requests.get("http://api.clicky.com/api/stats/4?site_id=32020&type=visitors,bounce-rate,time-average&date=today&output=json")
    api_request_7day = requests.get("http://api.clicky.com/api/stats/4?site_id=32020&type=visitors,bounce-rate,time-average&date=last-7-days&output=json")
    api_request_month = requests.get("http://api.clicky.com/api/stats/4?site_id=32020&type=visitors,bounce-rate,time-average&date=this-month&output=json")
    api_request_lastmonth = requests.get("http://api.clicky.com/api/stats/4?site_id=32020&type=visitors,bounce-rate,time-average&date=last-month&output=json")
    api_request_lastweek = requests.get("http://api.clicky.com/api/stats/4?site_id=32020&type=visitors,bounce-rate,time-average&date=last-week&output=json")
    api_request_links_domains = requests.get("http://api.clicky.com/api/stats/4?site_id=32020&type=links-domains&date=last-14-days&output=json")

    try:
        api_today = json.loads(api_request_today.content)
        api_7day = json.loads(api_request_7day.content)
        api_month = json.loads(api_request_month.content)
        api_lastmonth = json.loads(api_request_lastmonth.content)
        api_lastweek = json.loads(api_request_lastweek.content)
        api_links_domains = json.loads(api_request_links_domains.content)

    except Exception as e:
        api = "Error..."

    r = requests.get("http://api.clicky.com/api/stats/4?site_id=32020&type=visitors&date=this-month&output=json")
    month_json = r.json()
    month_string = json.dumps(month_json[0]['dates'][0]['date'], indent=2)
    month_string = month_string[-11:]
    month_string = month_string[:-1]
    #month_string = datetime.strptime(month_string, '%Y/%m/%d')
    month_int = ast.literal_eval(json.dumps(month_json[0]['dates'][0]['items'][0]['value']))


    '''
    if Visitor.objects.filter(visitor_month=month_string).exists():
        t = Visitor.objects.get(visitor_month=month_string)
        t.visitor_number = month_int
        t.save()

    else:
        t = Visitor(visitor_month=month_string, visitor_number=month_int)
        t.save()
        '''

    x = datetime.now()
    today = x.strftime('%d' + '.' + ' %b' + ' %Y')
    context = {
        'month_int':month_int,
        'month_string': month_string, 
        'api_today': api_today, 
        'api_7day': api_7day, 
        'api_month': api_month, 
        'api_lastmonth': api_lastmonth, 
        'api_lastweek': api_lastweek, 
        'api_links_domains': api_links_domains, 
        'time': today }
    return render(request, 'cms-kennzahlen-webseite.html', context)

@staff_member_required
def cms_marken(request):
    marken = Marke.objects.all().order_by('-id')            

    context = {
            'marken': marken,       
             }
    return render(request, 'cms-marken.html', context)

@staff_member_required
def cms_marke_erfassen(request):
    if request.method == "POST":
        form = MarkeChangeForm(request.POST or None, request.FILES or None)
        if form.is_valid():
            form.save()
            return redirect('store:cms_marken')

        else:
            messages.error(request, "Error")

    else: 
        form = MarkeChangeForm()

    context = {
        'form': form,
                }
    return render(request, 'cms-marke-erfassen.html', context)


@staff_member_required
def cms_marke_bearbeiten(request, pk):
    marke = get_object_or_404(Marke, pk=pk)
    if request.method == "POST":
        form = MarkeChangeForm(request.POST or None, instance=marke)
        if form.is_valid():
            form.save()
            return redirect('store:cms_marken')

        else:
            messages.error(request, "Error")
            
    else:
        form = MarkeChangeForm(request.POST or None, request.FILES or None, instance=marke)
        context = {
            'form': form,
            'marke': marke,
             }
        return render(request, 'cms-marke-bearbeiten.html', context)

@staff_member_required
def cms_marke_löschen(request, pk):
    eintrag = get_object_or_404(Marke, pk=pk)
    eintrag.delete()
    messages.info(request, "Die Marke wurde gelöscht.")
    return redirect("store:cms_marken") 

@staff_member_required
def cms_product_marke_overview(request, pk):
    product = get_object_or_404(Item, id=pk)

    product_marke = product.item_marken.all

    context = {
            'product': product,
            'product_marke' : product_marke,        
             }
    return render(request, 'cms-produkt-marken.html', context)

@staff_member_required
def cms_product_marke_erfassen(request, pk):

    product = get_object_or_404(Item, id=pk)

    if request.method == 'POST':
        form = ProductMarkeLinkForm(request.POST or None)
        if form.is_valid():
            cleaned_data = form.cleaned_data
            item_marke = cleaned_data['item_marke']
            marke = get_object_or_404(Marke, name=item_marke)
            marke.item.add(product)
            marke.save()
            return redirect('store:cms_product_marke_overview', pk=product.id)
    else:
            form = ProductMarkeLinkForm()

    context = {

        'product' : product,
        'form': form,
                }

    return render(request, 'cms-produkt-marke-erfassen.html', context)


@staff_member_required
def cms_product_marke_löschen(request, pkk, pk):
    product = get_object_or_404(Item, id=pkk)
    item_marke = get_object_or_404(Marke, id=pk)
    marke = get_object_or_404(Marke, name=item_marke)
    marke.item.remove(product)
    marke.save()
            
    return redirect('store:cms_product_marke_overview', pk=product.id)
    

    context = {

        'product' : product,
        'form': form,
                }

    return render(request, 'cms-produkt-marke-erfassen.html', context)



def lieferanten(request):
    search_query = request.GET.get('search', '')
    if search_query:
        lieferanten = Lieferanten.objects.filter(
            Q(name__icontains=search_query) |
            Q(adresse__icontains=search_query) |
            Q(ort__icontains=search_query) |
            Q(plz__icontains=search_query)
        ).order_by('-id')
    else:
        lieferanten = Lieferanten.objects.all().order_by('-id')

    context = {
        'lieferanten': lieferanten,
    }
    return render(request, 'crm/lieferanten.html', context)

def lieferant_create(request):
    if request.method == 'POST':
        form = LieferantenForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('store:lieferanten')
    else:
        form = LieferantenForm()

    return render(request, 'crm/lieferanten-erstellen.html', {'form': form})

def lieferant_edit(request, pk):
    lieferant = get_object_or_404(Lieferanten, pk=pk)
    if request.method == 'POST':
        form = LieferantenForm(request.POST, instance=lieferant)
        if form.is_valid():
            form.save()
            return redirect('store:lieferanten')
    else:
        form = LieferantenForm(instance=lieferant)

    return render(request, 'crm/lieferanten-erstellen.html', {'form': form})

@staff_member_required
def lieferant_delete(request, pk):
    eintrag = get_object_or_404(Lieferanten, pk=pk)
    eintrag.delete()
    messages.info(request, "Der Lieferant wurde gelöscht.")
    return redirect("store:lieferanten")



#webshop kunde

def cms_kunden(request):
    search_query = request.GET.get('search', '')
    
    # Retrieve User and Kunde separately, including search functionality
    if search_query:
        users = User.objects.filter(
            Q(profile__firmenname__icontains=search_query) | 
            Q(profile__interne_nummer__icontains=search_query)
        ).order_by('-id')
        
        kunden = Kunde.objects.filter(
            Q(firmenname__icontains=search_query) | 
            Q(interne_nummer__icontains=search_query)
        ).order_by('-id')
    else:
        users = User.objects.all().order_by('-id')
        kunden = Kunde.objects.all().order_by('-id')
    
    context = {
        'users': users,
        'kunden': kunden,
    }
    
    return render(request, 'cms-kunden.html', context)


def cms_kunden_erfassen(request):
    if request.method == "POST":
        kunde_form = KundeEditAdvancedForm(request.POST)
        address_form = AddressForm(request.POST)

        if kunde_form.is_valid() and address_form.is_valid():
            # Save Kunde and Address as needed
            new_user = User.objects.create_user(  # Example user creation
                username=kunde_form.cleaned_data['firmenname'],
                password=User.objects.make_random_password(),
                email=kunde_form.cleaned_data.get('email', '')
            )

            kunde = kunde_form.save(commit=False)
            kunde.user = new_user
            kunde.save()

            address = address_form.save(commit=False)
            address.kunde = kunde
            address.save()

            return redirect('store:cms_kunden')
        else:
            messages.error(request, "Fehler beim Erstellen des Kunden.")

    else:
        kunde_form = KundeEditAdvancedForm()
        address_form = AddressForm()

    return render(request, 'cms-kunden-erfassen.html', {
        'kunde_form': kunde_form,
        'address_form': address_form
    })


@staff_member_required
def cms_user_bearbeiten(request, pk):
    kunde = get_object_or_404(User, pk=pk)
    if request.method == "POST":
        form = KundeEditForm(request.POST or None, request.FILES or None, instance=kunde)
        if form.is_valid():
            form.save()
            return redirect('store:cms_kunden')

        else:
            messages.error(request, "Error")
            
    else:
        form = KundeEditForm(request.POST or None, instance=kunde)
        context = {
            'form': form,
            'kunde': kunde,
             }
        return render(request, 'cms-user-bearbeiten.html', context)


@staff_member_required
def cms_kunde_bearbeiten(request, pk):
    kunde = get_object_or_404(Kunde, pk=pk)
    try:
        address = CRMAddress.objects.get(kunde=kunde)  # Get the related address
    except CRMAddress.DoesNotExist:
        address = None  # Handle the case where the address does not exist

    if request.method == 'POST':
        # Initialize the forms with POST data and existing instances
        kunde_form = CRMKundeForm(request.POST, instance=kunde)
        kunde_form2 = CRMKundeRestForm(request.POST, instance=kunde)
        address_form = CRMAddressForm(request.POST, instance=address)

        if kunde_form.is_valid() and kunde_form2.is_valid() and address_form.is_valid():
            # Save kunde_form first
            kunde = kunde_form.save()

            # Save kunde_form2 with the updated kunde instance
            kunde_form2.instance = kunde
            kunde_form2.save()

            # Save the updated address
            address = address_form.save(commit=False)
            address.kunde = kunde  # Ensure the address is associated with the correct Kunde
            address.address_type = 'R'  # Keep the same address type or update as needed
            address.save()

            return redirect('store:cms_kunden')

    else:
        # Prepopulate the forms with the existing Kunde and Address data
        kunde_form = CRMKundeForm(instance=kunde)
        kunde_form2 = CRMKundeRestForm(instance=kunde)
        address_form = CRMAddressForm(instance=address)

    return render(request, 'cms-kunden-bearbeiten.html', {
        'kunde_form': kunde_form,
        'kunde_form2': kunde_form2,
        'address_form': address_form,
    })



@staff_member_required
def cms_kundenadresse_bearbeiten(request, pk):
    kunde = get_object_or_404(User, pk=pk)
    address = Address.objects.get(user=kunde)
    if request.method == "POST":
        form = AddressForm(request.POST or None, instance=address)
        if form.is_valid():
            form.save()
            return redirect('store:cms_kunden')

        else:
            messages.error(request, "Error")
            
    else:
        form = AddressForm(request.POST or None, instance=address)
        context = {
            'form': form,
            'kunde': kunde,
             }
        return render(request, 'cms-kundenadresse-bearbeiten.html', context)


@staff_member_required
def cms_kunde_löschen(request, pk):
    eintrag = get_object_or_404(User, pk=pk)
    eintrag.delete()
    messages.info(request, "Der Kunde wurde gelöscht.")
    return redirect("store:cms_kunden") 


@staff_member_required
def cms_produkte(request, first_cat):
    category = Category.objects.all()
    filter_query = request.GET.get('category', '')
    search_query = request.GET.get('search', '')
    first_cat_new = ''
    current_cat = 'standard'

    if search_query:
        produkte = Item.objects.filter(Q(titel__icontains=search_query) | Q(artikelnr__icontains=search_query) | Q(kategorie__name__icontains=search_query) | Q(subkategorie__sub_name__icontains=search_query))
        
        if filter_query:
            produkte = produkte.filter(kategorie__name=filter_query)
            current_cat_query = get_object_or_404(Category, name=filter_query)
            current_cat = current_cat_query.name 
            first_cat_new = current_cat
        else:
            pass
    else:
        if filter_query:
            produkte = Item.objects.filter(kategorie__name=filter_query)
            current_cat_query = get_object_or_404(Category, name=filter_query)
            current_cat = current_cat_query.name 
            first_cat_new = current_cat
        else: 
            produkte = Item.objects.filter(kategorie__name=first_cat)
            
    if current_cat == 'standard':
        current_cat = first_cat
    else:
        pass


    context = {
        'category': category,
        'produkte': produkte,
        'first_cat_new' : first_cat_new,
        'current_cat' : current_cat
     }
    return render(request, 'cms-produkte.html', context)


@staff_member_required
def product_cms_create(request, cat):
    if request.method == "POST":
        form = ProduktCreateForm(request.POST or None, request.FILES or None)
        if form.is_valid():
            result = form.save(commit=False)
            result.kategorie = get_object_or_404(Category, name=cat)
            result.save()
            return redirect('store:cms_produkte', first_cat=cat)
        else:
            messages.error(request, "Error")
    else:   
        initial_data = {
                'kategorie': cat,
                
            }
        form = ProduktCreateForm(initial=initial_data)

    context = {
        'form': form,
        'cat' : cat,
                }
    return render(request, 'cms-produkte-erfassen.html', context)


@staff_member_required
def product_cms_edit(request, pk, current_cat):
    item = get_object_or_404(Item, pk=pk)
    
    if request.method == "POST":
        form = ProduktEditForm(request.POST or None, request.FILES or None, instance=item)
        if form.is_valid():
            form.save()
            return redirect('store:cms_produkte', first_cat=current_cat)

        else:
            messages.error(request, "Error")
    else: 
        form = ProduktEditForm(request.POST or None, request.FILES or None, instance=item)

    context = {
        'form': form,
        'item': item,
        'current_cat': current_cat
                }
    return render(request, 'cms-produkte-bearbeiten.html', context)

@staff_member_required
def cms_remove_product(request, pk, cat):
    eintrag = get_object_or_404(Item, pk=pk)
    eintrag.delete()
    messages.info(request, "Der Eintrag wurde gelöscht.")
    return redirect("store:cms_produkte", first_cat=cat)    





@staff_member_required
def cms_versandkosten(request):
    shippingcost = ShippingCost.objects.all()
    context = {
        'object' : shippingcost,
    }
    return render(request, 'cms-versandkosten.html', context)

@staff_member_required
def cms_versandkosten_erfassen(request):
    if request.method == "POST":
        form = VersandkostenCreateForm(request.POST or None, request.FILES or None)
        if form.is_valid():
            form.save()
            return redirect('store:cms_versandkosten')

        else:
            messages.error(request, "Error")

    else: 
        form = form = VersandkostenCreateForm()

    context = {
        'form': form,
                }
    return render(request, 'cms-versandkosten-erfassen.html', context)

@staff_member_required
def cms_remove_versandkosten(request, pk):
    eintrag = get_object_or_404(ShippingCost, pk=pk)
    eintrag.delete()
    messages.info(request, "Der Eintrag wurde gelöscht.")
    return redirect("store:cms_versandkosten")  

@staff_member_required
def cms_statistik_produkte(request):
    labels = []
    data = []


    queryset = Item.objects.values(
            'titel'
        ).annotate(
            total_sales=Coalesce(Sum('orderitem__quantity'), 0)
        ).order_by('-total_sales')[:5]

    for entry in queryset:
        labels.append(entry['titel'])
        data.append(entry['total_sales'])

    context = {
        'queryset': queryset,
        'labels': labels,
        'data': data,       
    }

    return render(request, 'cms-statistik-produkte.html', context)

@staff_member_required
def cms_benutzerdaten(request):
    context = { }
    return render(request, 'cms-benutzerdaten.html', context)


@login_required
def login_user(request):
    if request.method == "POST":
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            messages.success(request, ('Erfolgreich angemeldet.'))
            return redirect ('store:cms')
        else:
            messages.success(request, ('Fehler - bitte überprüfen Sie Ihre Angaben.'))
            return redirect ('store:login_user')
    else:
        context = {}
        return render(request, 'cms-login.html', context )

@login_required
def logout_user(request):
    logout(request)
    messages.success(request, ('Sie wurden erfolgreich abgemeldet.'))
    return redirect('store:login_user')



@staff_member_required
def bestellung_erfassen_view(request):
    if request.method == 'POST':
        form = BestellungForm(request.POST)
        kunden_nr = request.POST.get('kunden_nr')

        # Check if Kunde exists
        kunde = Kunde.objects.filter(interne_nummer=kunden_nr).first()
        
        if not kunde:
            messages.error(request, f"Kunde mit Nummer '{kunden_nr}' wurde nicht gefunden.")
        elif not Kunde.objects.filter(interne_nummer=kunden_nr).exists():
            messages.error(request, f"Kunde mit Nummer '{kunden_nr}' wurde nicht gefunden.")
        
        # Check if the Kunde has at least one related 'Elemente' record
        elif not kunde.kunden_elemente.exists():
            messages.error(request, "Dieser Kunde hat noch keine Elemente. Eine Bestellung kann nur erfasst werden, wenn mindestens ein Element vorhanden ist.")
        
        # If form is valid, proceed with saving the Bestellung
        elif form.is_valid():
            bestellung = form.save(commit=False)
            bestellung.kunden_nr = kunden_nr  # Kunden-Nr. manuell setzen
            bestellung.save()  # Speichern der Bestellung
            messages.success(request, "Die Bestellung wurde erfolgreich erfasst.")
            return redirect('store:elemente_bestellungen')  # Redirect after success
        else:
            messages.error(request, "Bitte überprüfen Sie Ihre Eingaben.")
    else:
        form = BestellungForm()

    # Fetch list of customers excluding customers with no 'interne_nummer'
    kunden_liste = Kunde.objects.exclude(interne_nummer__isnull=True)

    return render(request, 'crm/bestellung_form.html', {
        'bestellung_form': form,
        'kunden_liste': kunden_liste,
    })